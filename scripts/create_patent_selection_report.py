# -*- coding: utf-8 -*-
"""Create the A2 humanoid patent screening criteria report as a polished DOCX."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "outputs" / "휴머노이드_유효특허_선별기준_및_판정사례_v1.docx"

# standard_business_brief preset with a named Korean typography override.
FONT = "Malgun Gothic"
NAVY = "183B56"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "24313D"
MUTED = "5F6B76"
LIGHT_BLUE = "EAF2F8"
PALE_BLUE = "F4F8FB"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D7DEE5"
WHITE = "FFFFFF"
GREEN = "2E6B4F"
PALE_GREEN = "EAF5EF"
GOLD = "8A6514"
PALE_GOLD = "FFF7E3"
RED = "9B1C1C"
PALE_RED = "FCEEEE"

PAGE_WIDTH_DXA = 12240  # Letter, 8.5 in
PAGE_HEIGHT_DXA = 15840  # Letter, 11 in
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_TOP_BOTTOM_DXA = 80
CELL_SIDE_DXA = 120


def set_run_font(run, size=None, bold=None, italic=None, color=INK):
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_style_font(style, size, color=INK, bold=False):
    style.font.name = FONT
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), FONT)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_margins(cell, top=CELL_TOP_BOTTOM_DXA, start=CELL_SIDE_DXA,
                     bottom=CELL_TOP_BOTTOM_DXA, end=CELL_SIDE_DXA):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_borders(cell, color=MID_GRAY, size=4):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = tc_borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            tc_borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths):
    assert sum(widths) == CONTENT_WIDTH_DXA
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[idx]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            set_cell_borders(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_paragraph_shading(paragraph, fill, border_color=None):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    if border_color:
        p_bdr = p_pr.find(qn("w:pBdr"))
        if p_bdr is None:
            p_bdr = OxmlElement("w:pBdr")
            p_pr.append(p_bdr)
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "20")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), border_color)
        p_bdr.append(left)


def keep_with_next(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    node = OxmlElement("w:keepNext")
    p_pr.append(node)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:cantSplit")
    tr_pr.append(node)


def add_field(run, instruction):
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_separate, display, fld_char_end])


def add_numbering_definition(doc, num_fmt, text, left=720, hanging=360):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(x.get(qn("w:abstractNumId")))
        for x in numbering.findall(qn("w:abstractNum"))
        if x.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(x.get(qn("w:numId")))
        for x in numbering.findall(qn("w:num"))
        if x.get(qn("w:numId")) is not None
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(left))
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(left))
    ind.set(qn("w:hanging"), str(hanging))
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, ind, spacing])
    lvl.extend([start, fmt, lvl_text, suff, p_pr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])


def add_para(doc, text="", style=None, after=6, before=0, line=1.1,
             align=WD_ALIGN_PARAGRAPH.LEFT, bold_prefix=None, color=INK,
             keep=False):
    paragraph = doc.add_paragraph(style=style)
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    if keep:
        keep_with_next(paragraph)
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        set_run_font(first, bold=True, color=color)
        rest = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(rest, color=color)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, color=color)
    return paragraph


def add_bullet(doc, text, bullet_num_id, bold_prefix=None, after=5):
    paragraph = add_para(doc, text, after=after, line=1.167, bold_prefix=bold_prefix)
    apply_numbering(paragraph, bullet_num_id)
    return paragraph


def add_step(doc, text, decimal_num_id, bold_prefix=None, after=6):
    paragraph = add_para(doc, text, after=after, line=1.167, bold_prefix=bold_prefix)
    apply_numbering(paragraph, decimal_num_id)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    keep_with_next(paragraph)
    return paragraph


def add_kicker(doc, text):
    p = add_para(doc, text, after=5, before=0, color=BLUE, keep=True)
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(10)
    return p


def add_callout(doc, label, text, fill=LIGHT_BLUE, accent=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.06)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.15
    set_paragraph_shading(p, fill, accent)
    r1 = p.add_run(f"{label}  ")
    set_run_font(r1, size=10.5, bold=True, color=accent)
    r2 = p.add_run(text)
    set_run_font(r2, size=10.5, color=INK)
    return p


def add_label_line(doc, label, text, after=3):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.05
    r1 = p.add_run(label)
    set_run_font(r1, size=10.5, bold=True, color=NAVY)
    r2 = p.add_run(text)
    set_run_font(r2, size=10.5, color=INK)
    return p


def set_cell_text(cell, text, bold=False, color=INK, size=9.3,
                  align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.12
    run = p.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color)


def add_table(doc, headers, rows, widths, font_size=9.3, header_fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_text(cell, header, bold=True, color=NAVY, size=9.2,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, header_fill)
    set_repeat_table_header(table.rows[0])

    for row_values in rows:
        row = table.add_row()
        prevent_row_split(row)
        for idx, value in enumerate(row_values):
            align = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(row.cells[idx], value, size=font_size, align=align)
    set_table_geometry(table, widths)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1
    return table


def add_source_note(doc, text):
    p = add_para(doc, text, before=4, after=8, line=1.05, color=MUTED)
    p.runs[0].font.size = Pt(8.5)
    return p


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    set_style_font(normal, 11, INK, False)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        set_style_font(style, size, color, True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.05
        style.paragraph_format.keep_with_next = True

    if "Case Label" not in doc.styles:
        case_style = doc.styles.add_style("Case Label", WD_STYLE_TYPE.PARAGRAPH)
    else:
        case_style = doc.styles["Case Label"]
    set_style_font(case_style, 10.5, DARK_BLUE, True)
    case_style.paragraph_format.space_before = Pt(7)
    case_style.paragraph_format.space_after = Pt(3)
    case_style.paragraph_format.keep_with_next = True


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    hr = hp.add_run("휴머노이드 로봇 유효특허 선별 기준  |  A2 정답셋")
    set_run_font(hr, size=8.5, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    fr = fp.add_run("선별 기준 v1  ·  2026.07.24  ·  ")
    set_run_font(fr, size=8.5, color=MUTED)
    page_run = fp.add_run()
    set_run_font(page_run, size=8.5, color=MUTED)
    add_field(page_run, "PAGE")


def build_document():
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    bullet_num_id = add_numbering_definition(doc, "bullet", "•", left=720, hanging=360)
    decimal_num_id = add_numbering_definition(doc, "decimal", "%1.", left=720, hanging=360)

    # Opening block: memo_masthead adapted for a formal technical report.
    add_kicker(doc, "A2 정답셋 구축 · 유효특허 선별 기준 v1")
    title = add_para(
        doc,
        "휴머노이드 로봇 유효특허\n선별 기준 및 판정 사례 보고서",
        after=6,
        line=1.0,
        color=NAVY,
        keep=True,
    )
    for run in title.runs:
        set_run_font(run, size=23, bold=True, color=NAVY)
    subtitle = add_para(
        doc,
        "KIMM 가치사슬 3축과 확정 판정 규칙 12개에 기반한 포함·제외 체계",
        after=14,
        line=1.1,
        color=MUTED,
        keep=True,
    )
    subtitle.runs[0].font.size = Pt(12.5)

    add_label_line(doc, "문서 목적  ", "특허 데이터 정제 과정에서 적용한 기술 관련성 기준과 대표 판정 사례를 보고 가능한 형식으로 명문화")
    add_label_line(doc, "분석 범위  ", "미국 특허, 출원일 2015.01.01~2025.12.31, 원자료 3,757건 → 패밀리 대표 3,323건")
    add_label_line(doc, "판정 체계  ", "T1(직접 핵심) / T2(이전가능 코어) / E(제외), T1+T2를 분석상 유효특허로 채택")
    add_label_line(doc, "기준 상태  ", "2026.07.24 확정 규칙 v1")

    add_callout(
        doc,
        "핵심 결론",
        "출원인이나 키워드만으로 포함 여부를 결정하지 않았다. 발명의 청구 대상과 기술적 핵심을 우선 확인하고, "
        "휴머노이드 상용화 가치사슬에 직접 해당하거나 휴머노이드로 이전 가능한 코어 기술만 유효특허로 포함했다.",
    )

    add_heading(doc, "1. 선별의 목적과 기본 원칙", 1)
    add_para(
        doc,
        "본 선별의 목적은 ‘로봇’이라는 표현이 들어간 특허를 넓게 수집하는 것이 아니라, 휴머노이드 로봇의 상용화와 "
        "기술 경쟁력 분석에 실제로 필요한 특허 집합을 구성하는 데 있다. 따라서 범용 산업용 로봇, 수술로봇, 청소로봇, "
        "물류 자동화 등 인접 분야의 특허는 분야명 자체가 아니라 발명의 기술적 본체와 휴머노이드 이전가능성을 기준으로 "
        "개별 판정하였다.",
    )
    add_callout(
        doc,
        "용어 주의",
        "이 문서에서 ‘유효특허’는 분석 목적상 기술 관련성이 인정되어 최종 분석집합에 포함된 특허를 의미한다. "
        "특허권의 존속 여부, 무효 가능성 또는 법률적 권리 유효성을 판단한 것은 아니다.",
        fill=PALE_GOLD,
        accent=GOLD,
    )

    add_heading(doc, "2. 선별 결과를 구성하는 세 가지 라벨", 1)
    add_para(
        doc,
        "판정 결과는 관련성 점수의 임의 컷오프가 아니라, 기술의 직접성·이전가능성·용도 전용성을 기준으로 다음 세 라벨 중 "
        "하나로 확정하였다. 최종 분석에서는 T1과 T2를 유효특허로 포함하고 E는 제외한다.",
    )
    add_table(
        doc,
        ["라벨", "의미", "포함되는 경우", "최종 처리"],
        [
            (
                "T1",
                "직접 핵심",
                "휴머노이드·인간형·이족보행·전신제어를 직접 대상으로 하거나 완제품 휴머노이드의 핵심 서브시스템인 경우",
                "유효특허 포함",
            ),
            (
                "T2",
                "이전가능 코어",
                "휴머노이드 전용은 아니지만 H/W·지능·데이터 축에서 휴머노이드로의 기술 이전성이 명시적이고 핵심적인 경우",
                "유효특허 포함",
            ),
            (
                "E",
                "제외",
                "특정 공정·설비·제품·서비스에 묶이거나 키워드·응용례만 휴머노이드와 접점이 있는 경우",
                "분석집합 제외",
            ),
        ],
        [850, 1450, 5180, 1880],
        font_size=9.2,
    )
    add_source_note(doc, "판정 단위: 동일 발명의 계속출원·패밀리는 대표 1건으로 집계한다.")

    add_page_break(doc)
    add_heading(doc, "3. 기술 관련성의 기준: KIMM 가치사슬 3축", 1)
    add_para(
        doc,
        "포함 가능 기술의 범위는 KIMM 보고서가 제시한 휴머노이드 가치사슬의 세 축, 즉 H/W·지능·데이터에 앵커링하였다. "
        "다만 세 축 중 하나에 이름상 속한다는 사실만으로 자동 포함하지 않고, 해당 특허가 휴머노이드의 기능 또는 상용화 "
        "병목을 실제로 해결하는지를 추가로 확인하였다.",
    )
    add_table(
        doc,
        ["기술 축", "핵심 구성", "대표 세부기술", "포함 판단 질문"],
        [
            (
                "H/W",
                "신체·구동·감지·전원",
                "QDD·로터리/리니어 액추에이터, 감속기, 롤러스크류, 모션제어, 힘/토크·촉각센서, 로봇손, 경량 소재, 배터리·열관리",
                "휴머노이드의 조작·보행·안전·가동시간 병목을 직접 해결하는가?",
            ),
            (
                "지능",
                "인지·학습·행동 생성",
                "VLA, E2E, RFM, SLAM, 강화학습·모방학습, 비전·행동계획, AI칩·온디바이스 제어",
                "일반 AI가 아니라 로봇의 지각·조작·보행 행동을 생성하거나 개선하는가?",
            ),
            (
                "데이터",
                "학습 데이터 생산·이식·운영",
                "텔레오퍼레이션, 시연 데이터, Sim-to-Real, 도메인 무작위화, 플릿 운영, SDK·미들웨어, 데이터·사이버 보안",
                "실세계 데이터 부족이나 시뮬레이션-현실 격차 등 상용화 병목을 해결하는가?",
            ),
        ],
        [1000, 1700, 4000, 2660],
        font_size=9.05,
        header_fill=LIGHT_BLUE,
    )
    add_callout(
        doc,
        "적용 원칙",
        "3축은 ‘무엇을 볼 것인가’를 정하는 기술 범위이고, T1/T2/E와 12개 규칙은 ‘그 기술을 최종적으로 포함할 것인가’를 "
        "결정하는 판정 체계다. 예를 들어 배터리나 AI칩이라는 이유만으로 포함하지 않으며, 휴머노이드 적용과 해결 기능이 "
        "특허의 핵심으로 확인되어야 한다.",
    )

    add_heading(doc, "4. 실제 선별 절차", 1)
    add_para(
        doc,
        "선별은 넓게 후보를 찾은 뒤 기술적 본체를 확인하는 방식으로 수행하였다. 키워드와 출원인은 후보 발견 및 검토 순서 "
        "결정에만 사용하고, 최종 라벨은 특허 단위의 내용 판정으로 확정하였다.",
    )
    steps = [
        ("패밀리 중복 제거", "동일 발명의 계속출원·패밀리를 대표 1건으로 통합하여 건수 왜곡을 방지했다."),
        ("출원인 정규화와 후보 탐색", "출원인 표기를 정규화하고 제목·초록·대표청구항·AI 요약에서 관련 신호를 추출했다."),
        ("직접 휴머노이드 여부 확인", "휴머노이드·이족보행·전신제어 또는 완제품 플랫폼을 직접 청구하면 T1 후보로 분류했다."),
        ("3축 및 이전가능성 검토", "직접 휴머노이드가 아니더라도 H/W·지능·데이터 축의 코어 기술이며 이전성이 명시적이면 T2로 분류했다."),
        ("용도 전용성 및 예외 규칙 적용", "특정 공정·설비·제품에 묶인 경우 E로 분류하고, 4족·물류학습·산업용 회색지대 등 확정 규칙을 적용했다."),
        ("경계 사례 HITL 확인", "자동 규칙으로 확정하기 어려운 사례는 포함·제외 대비 사례와 함께 사람의 최종 검토 대상으로 남겼다."),
    ]
    for label, detail in steps:
        add_step(doc, f"{label}: {detail}", decimal_num_id, bold_prefix=f"{label}:")

    add_callout(
        doc,
        "판정 우선순위",
        "청구 대상·필수 한정요소 > 발명의 핵심 기능 > 다른 로봇으로의 이전가능성 > 응용례 > 출원인·키워드",
        fill=PALE_GREEN,
        accent=GREEN,
    )

    add_page_break(doc)
    add_heading(doc, "5. 포함 기준", 1)
    add_para(
        doc,
        "포함은 두 단계로 이루어진다. 먼저 휴머노이드 자체 또는 그 직접 서브시스템이면 T1으로 포함하고, 그 외에는 "
        "휴머노이드 상용화에 이전 가능한 코어 기술인지 확인하여 T2로 포함한다.",
    )

    add_heading(doc, "5.1 T1: 직접 핵심 특허", 2)
    add_bullet(
        doc,
        "휴머노이드·인간형·이족보행 로봇의 본체 구조, 전신제어, 균형·보행 기술",
        bullet_num_id,
    )
    add_bullet(
        doc,
        "완제품 휴머노이드 플랫폼의 몸통·팔·손·발·관절 및 안전 제어",
        bullet_num_id,
    )
    add_bullet(
        doc,
        "OEM 목록 밖 기업의 특허라도 휴머노이드 적용이 발명의 본체로 명시된 경우",
        bullet_num_id,
    )
    add_callout(
        doc,
        "대표 사례",
        "Figure AI US12605824(휴머노이드 전신·엄지 직접구동), Boston Dynamics US11911916(Atlas 계열 전신 균형), "
        "UBTech US11045945(휴머노이드 보행 충격 제어), Agility Robotics US12552606(Digit의 양손 조작), "
        "Xiaomi US12440966(이족 질량중심 예측 제어)는 T1으로 포함했다.",
    )

    add_heading(doc, "5.2 T2: 휴머노이드 이전가능 코어 특허", 2)
    add_para(
        doc,
        "다른 로봇 형태나 산업에서 출발한 기술이라도, 발명의 핵심이 휴머노이드의 조작·보행·학습·안전·데이터 병목에 "
        "직접 재사용될 수 있고 특정 공정에 종속되지 않으면 T2로 포함한다.",
    )
    add_bullet(
        doc,
        "조작·로봇손: 다지 핸드, 촉각·힘 센싱, 범용 파지 계획, 인핸드 조작",
        bullet_num_id,
        bold_prefix="조작·로봇손:",
    )
    add_bullet(
        doc,
        "구동·균형: QDD·SEA·경량 관절, 감속기, 4족·휠레그의 전신제어와 다리 기술",
        bullet_num_id,
        bold_prefix="구동·균형:",
    )
    add_bullet(
        doc,
        "지능·학습: VLA, 모방·강화학습, 범용 조작 스킬, 접촉·힘 기반 제어",
        bullet_num_id,
        bold_prefix="지능·학습:",
    )
    add_bullet(
        doc,
        "데이터·운영: 텔레오퍼레이션, 시연 데이터 수집, Sim-to-Real, 도메인 무작위화",
        bullet_num_id,
        bold_prefix="데이터·운영:",
    )
    add_bullet(
        doc,
        "인간협업 안전: 접촉력 검출, 안전 엔벨로프, 힘 제한, 프리드라이브 안전",
        bullet_num_id,
        bold_prefix="인간협업 안전:",
    )
    add_callout(
        doc,
        "대표 사례",
        "Samsung US11865714(QDD 이중감속 관절), Google US12528186(VLA 정책), Acumino US11822710(시연 데이터 수집), "
        "Stanford US12269156(햅틱 텔레옵), Veo Robotics US11254004(정지거리 기반 협업 안전)는 휴머노이드 "
        "상용화의 코어 기능으로 이전 가능하므로 T2에 포함했다.",
    )

    add_heading(doc, "5.3 기업·응용 분야보다 발명의 본체를 우선", 2)
    add_para(
        doc,
        "휴머노이드 기업의 특허라도 제품라인이 다르면 판정이 달라지고, 반대로 수술·외골격·물류 기업의 특허라도 청구 대상이 "
        "범용 코어 기술이면 포함될 수 있다. 따라서 기업 단위 일괄 포함·일괄 제외를 적용하지 않았다.",
    )
    add_table(
        doc,
        ["비교", "포함 사례", "제외·하향 사례", "판정 이유"],
        [
            (
                "동일 기업",
                "Boston Dynamics Atlas 전신균형 → T1",
                "Stretch 흡착 물류 패키징 → E/T2",
                "기업명이 아니라 제품라인과 청구 기능이 다름",
            ),
            (
                "OEM 밖 구제",
                "Honda ASIMO·Toyota T-HR3 명시 → T1",
                "일반 서비스 조정 SW → E",
                "휴머노이드가 발명의 본체인지 확인",
            ),
            (
                "전업 분야 구제",
                "수술기업의 범용 임피던스 제어 → T2",
                "수술 셋업·카테터 시술 전용 → E",
                "범용 코어와 용도 전용 기술을 분리",
            ),
        ],
        [1150, 2600, 2500, 3110],
        font_size=8.9,
    )

    add_page_break(doc)
    add_heading(doc, "6. 제외 기준", 1)
    add_para(
        doc,
        "다음 유형은 원칙적으로 E로 제외하였다. 다만 카테고리 이름만으로 일괄 제외하지 않고, 청구 대상이 범용 코어 기술인 "
        "경우에는 T1 또는 T2로 구제하였다.",
    )

    exclusion_items = [
        (
            "특정 산업공정·설비 전용",
            "용접·도장·SMT·웨이퍼·공작기계·컨베이어 등 특정 공정의 문제 해결이 청구항의 필수 구성인 경우",
            "FANUC US9902070(공작기계 냉각제 노즐), US12447615(용접 로봇 교시)",
        ),
        (
            "의료·수술 목적 전용",
            "수술 셋업, 카테터 시술, 의료 영상 안내, 수술 가상경계 등 의료 절차 자체에 종속된 경우",
            "Medivis US12521893, Corindus US12023807, Stryker US11648679",
        ),
        (
            "외골격·재활·의족 전용",
            "착용자의 보행 보조·재활·의족 기능이 발명의 목적이며 로봇 코어로의 이전이 청구되지 않은 경우",
            "Roam Robotics 외골격 패밀리, MIT US9975249(의족)",
        ),
        (
            "청소·잔디·시설관리 전용",
            "바닥청소기 주행·충전, 외벽청소, 잔디관리 등 전용 제품 기능인 경우",
            "iRobot·LG·Samsung 청소로봇 주행·충전 특허",
        ),
        (
            "AGV·물류 인프라 전용",
            "반송·플릿·컨베이어·팔레트 배치·창고 워크플로가 청구 대상인 경우",
            "Dexterity US12485544(팔레타이징 배치 학습), GM US11338450(공정 AGV)",
        ),
        (
            "완구·엔터테인먼트 장치",
            "완구성 동작, 식용·포옹 애니매트로닉, 가상화 SW 등 로봇 실체 기술과 거리가 있는 경우",
            "Petoi US11833688, US11865695(포옹 애니매트로닉)",
        ),
        (
            "키워드·응용례만 관련",
            "휴머노이드가 가능한 응용례로만 열거되거나, 유사 단어의 의미가 기술 영역과 다른 경우",
            "recycling≠재활, load balancing≠균형제어, pedestrian≠보행",
        ),
    ]
    for label, criterion, case in exclusion_items:
        p = doc.add_paragraph(style="Case Label")
        r = p.add_run(label)
        set_run_font(r, size=10.5, bold=True, color=DARK_BLUE)
        add_label_line(doc, "제외 조건  ", criterion, after=2)
        add_label_line(doc, "대표 사례  ", case, after=6)

    add_callout(
        doc,
        "중요 예외",
        "수술·외골격·청소·물류·완구라는 카테고리만으로 자동 제외하지 않는다. 예를 들어 Samsung US11865714는 "
        "의료 문맥에서 검색되었지만 발명의 본체가 범용 QDD 관절이므로 T2로 구제했고, Agility US12440980은 "
        "물류 응용이라도 Digit 휴머노이드의 양팔 조작이므로 T1으로 포함했다.",
        fill=PALE_GOLD,
        accent=GOLD,
    )

    add_page_break(doc)
    add_heading(doc, "7. 경계 사례에 대한 확정 정책", 1)
    add_para(
        doc,
        "표본 검토에서 판정 편차가 가장 컸던 네 영역은 아래와 같이 확정하였다. 이 정책은 유사 사례에 일관되게 적용한다.",
    )

    add_heading(doc, "7.1 4족·다족·휠레그 기술", 2)
    add_para(
        doc,
        "전신 동역학, 균형, 다리 액추에이터, 지형 대응 기술은 이족 휴머노이드로의 이전성을 인정하여 T2로 포함하고 "
        "‘4족’ 또는 해당 폼팩터를 주석으로 남긴다. 완구용 4족은 E로 제외한다.",
    )

    add_heading(doc, "7.2 물류 학습", 2)
    add_para(
        doc,
        "학습 알고리즘 사용 여부가 아니라 학습 대상이 무엇인지로 나눈다. 빈피킹·파지·접촉 조작과 같은 범용 로봇 스킬을 "
        "학습하면 T2, 팔레타이징·분류·배치 등 물류 워크플로를 학습하면 E로 판정한다.",
    )

    add_heading(doc, "7.3 산업용 로봇 회색지대", 2)
    add_para(
        doc,
        "힘 기반 직접교시, 외력 추정, 촉각센싱, 범용 파지, 시뮬레이션 기반 힘제어 파라미터 튜닝은 T2로 포함한다. "
        "교시 UI, 캘리브레이션, 방수, 공정 검증용 OLP 및 특정 설비에 묶인 기능은 E로 제외한다.",
    )

    add_heading(doc, "7.4 진공흡착·소프트 그리퍼", 2)
    add_para(
        doc,
        "진공흡착은 기본적으로 E이지만 촉각·힘 피드백 조정, 손가락 하이브리드, 파지품질 학습·물리모델의 세 경우는 T2로 "
        "포함한다. 소프트 그리퍼는 파지 원리와 기구 자체가 청구 대상일 때 T2로 포함하되, 시스템 패키징은 T2 약으로 "
        "보수적으로 기록한다.",
    )

    add_table(
        doc,
        ["쟁점", "포함 판정", "제외 판정", "구분 기준"],
        [
            (
                "4족·휠레그",
                "UBTech US12496711: 다족 하중균형 → T2",
                "Petoi US11833688: 완구 4족 → E",
                "전신제어·다리 기술의 이족 이전성",
            ),
            (
                "물류 학습",
                "FANUC/PFN US11780095: 빈피킹 학습 → T2",
                "Dexterity US12485544: 팔레타이징 배치 → E",
                "스킬 학습인가, 워크플로 학습인가",
            ),
            (
                "산업용 로봇",
                "FANUC US11938633: 힘제어 시뮬 튜닝 → T2",
                "FANUC US12447615: 용접 교시 → E",
                "범용 코어인가, 특정 공정 문제인가",
            ),
            (
                "진공흡착",
                "MUJIN US10532462: 접촉 피드백 조정 → T2",
                "단순 흡착·반송 장치 → E",
                "피드백·하이브리드·품질학습 예외",
            ),
        ],
        [1400, 2700, 2480, 2780],
        font_size=8.75,
        header_fill=LIGHT_BLUE,
    )

    add_page_break(doc)
    add_heading(doc, "8. 확정 판정 규칙 12개", 1)
    add_para(
        doc,
        "다음 12개 규칙은 사례 검토 결과를 재현 가능한 판정 기준으로 정리한 최종 규칙이다. 규칙 1~8은 기본 원칙이며, "
        "규칙 9~12는 마지막 검토에서 확정된 구체화 규칙이다.",
    )

    rules = [
        (
            "1. 용도 전용성 원칙",
            "청구항이 특정 용도·공정에 묶이면 E, 로봇 일반의 파지·힘제어·학습·기구 자체를 청구하면 T2로 판정한다. "
            "특정 공정·설비가 청구항의 필수 전제나 한정요소이면 E를 우선한다.",
        ),
        (
            "2. OEM 제품라인 3분할 원칙",
            "완제품 OEM도 기업 단위로 일괄 포함하지 않는다. 예를 들어 Atlas는 T1, Spot 계열은 T2, Stretch 물류 "
            "계열은 E 또는 T2로 제품라인별 분리한다.",
        ),
        (
            "3. 역방향 구제 원칙",
            "OEM 목록 밖 기업이라도 휴머노이드·이족보행·전신제어가 발명의 본체로 명시되면 T1로 구제한다.",
        ),
        (
            "4. 폼팩터 주석 원칙",
            "Pepper·PR2 같은 바퀴형 상반신 휴머노이드는 T1과 함께 ‘바퀴형’ 주석을 남긴다. 4족·휠레그는 T2, "
            "수중 휴머노이드는 T2 약으로 기록한다.",
        ),
        (
            "5. 진공흡착 3예외 원칙",
            "진공흡착은 기본 E다. 단, 촉각·힘 피드백 조정, 손가락 하이브리드, 파지품질 학습·물리모델의 세 경우는 T2로 포함한다.",
        ),
        (
            "6. 오매칭 예외 처리 원칙",
            "recycling과 재활, 세척 용이 설계와 청소 기술을 구분하고, 청소·조립·배송 등 단순 응용례 나열은 포함 근거로 사용하지 않는다.",
        ),
        (
            "7. 패밀리 중복 제거 원칙",
            "사실상 동일 발명의 계속출원·패밀리는 대표 1건만 남겨 출원 건수와 기술 집중도 왜곡을 방지한다.",
        ),
        (
            "8. 데이터 결함 보정 원칙",
            "등록번호 공란, 출원인 필드 오염, 관리번호 혼입 등 원천데이터 결함을 보정한 뒤 판정한다.",
        ),
        (
            "9. 오탐 배제 사전 원칙",
            "soft actuator가 그리퍼 손가락인 경우, 서버 load balancing, gravity balancer, pedestrian, workpiece/tool pose, "
            "recycling, 물체의 safe handling, OLP·은유적 simulation은 키워드 단계에서 예외 처리한다.",
        ),
        (
            "10. 4족·다족·휠레그 확정 원칙",
            "전신제어·다리 액추에이터·균형 기술은 이족 이전성을 인정하여 T2로 포함하고 폼팩터 주석을 남긴다. 완구 4족만 E로 제외한다.",
        ),
        (
            "11. 물류 학습 분할 원칙",
            "빈피킹·파지·접촉조작 등 범용 스킬 학습은 T2, 팔레타이징·분류·배치 등 물류 워크플로 학습은 E로 판정한다.",
        ),
        (
            "12. 출원인 화이트리스트 가속 원칙",
            "주요 휴머노이드 플레이어는 우선 검토 클러스터로 배정하되 자동 포함하지 않는다. 확정 클러스터 배정 후에도 "
            "제품라인 3분할과 특허 단위 내용 판정을 적용한다.",
        ),
    ]
    for title_text, detail in rules:
        add_heading(doc, title_text, 2)
        add_para(doc, detail, after=5, line=1.1)

    add_callout(
        doc,
        "중복의 의미",
        "규칙 4와 10, 규칙 6과 9는 각각 일반원칙과 최종 구체화 규칙의 관계다. 서로 다른 기준이 충돌하는 것이 아니라, "
        "초기 원칙을 실제 자동·수동 판정에 적용할 수 있도록 세부 조건을 확정한 것이다.",
        fill=PALE_GOLD,
        accent=GOLD,
    )

    add_page_break(doc)
    add_heading(doc, "9. 대표 판정 사례", 1)
    add_para(
        doc,
        "아래 사례는 규칙을 적용할 때 참조할 수 있는 대표 선례다. 사례표는 판정 논리를 설명하기 위한 것이며 전체 특허의 "
        "전수 라벨 목록을 의미하지 않는다.",
    )
    add_table(
        doc,
        ["판정", "등록번호·출원인", "기술 내용", "판정 근거"],
        [
            ("T1", "US12605824 · Figure AI", "휴머노이드 전신·엄지 직접구동", "완제품 휴머노이드 구조를 직접 청구"),
            ("T1", "US11911916 · Boston Dynamics", "Atlas 계열 피치·롤 전신 균형", "휴머노이드 전신 균형 제어"),
            ("T1", "US11045945 · UBTech", "보행 단·양각 전환 충격 제어", "휴머노이드 이족보행 명시"),
            ("T1", "US12552606 · Agility Robotics", "Digit의 양손·비파지 조작", "물류 응용이어도 주체가 이족 휴머노이드"),
            ("T1", "US12179349 · Toyota", "T-HR3 계열 허리 모션 데이터 생성", "OEM 목록 밖 역방향 구제"),
            ("T2", "US11865714 · Samsung", "QDD 이중감속 준직접구동 관절", "의료 문맥 오매칭이나 범용 관절 코어"),
            ("T2", "US12172297 · Sanctuary", "힘줄·실리콘 스킨 다지 손", "휴머노이드 조작 병목인 로봇손"),
            ("T2", "US11524414 · Toyota", "로봇 손·팔용 6축 촉각센서", "힘·촉각 기반 조작 코어"),
            ("T2", "US12318935 · NVIDIA", "신경 가치함수 기반 실시간 파지", "범용 조작 스킬 학습"),
            ("T2", "US12528186 · Google", "자연어+이미지 목표조건 정책", "로봇 VLA 직계 기술"),
            ("T2", "US11822710 · Acumino", "착용형 인간 시연 데이터 수집", "외골격 형태이나 목적은 로봇 학습 데이터"),
            ("T2", "US11458630 · X Development", "접촉 시뮬레이션 기반 reality gap 완화", "Sim-to-Real 핵심"),
            ("T2", "US11254004 · Veo Robotics", "정지시간·거리 자동 추정", "인간협업 안전 핵심"),
            ("T2", "US12496711 · UBTech", "다족 하중균형", "4족·다족 전신제어의 이족 이전성"),
            ("E", "US9902070 · FANUC", "공작기계 냉각제 노즐 조정", "특정 공작기계 공정 전용"),
            ("E", "US12447615 · FANUC", "용접 로봇 교시", "용접 공정과 교시 기능에 종속"),
            ("E", "US12023807 · Corindus", "카테터 시술 로봇", "의료 시술 전용"),
            ("E", "US9975249 · MIT", "의족 신경근 모델", "착용 보조 목적 전용"),
            ("E", "US12485544 · Dexterity", "팔레타이징 배치 학습", "조작 스킬이 아닌 물류 워크플로 학습"),
            ("E", "US11833688 · Petoi", "완구용 4족 로봇", "4족 기술 예외 중 완구 전용"),
            ("E", "US9302393 · 개인", "휴머노이드 형태의 음성 AI 시스템", "휴머노이드 실체 기술이 아님"),
        ],
        [720, 2150, 3160, 3330],
        font_size=8.45,
        header_fill=LIGHT_BLUE,
    )
    add_source_note(doc, "주: 경계 사례는 동일 규칙 아래에서도 청구항의 필수 한정요소에 따라 최종 판정이 달라질 수 있다.")

    add_page_break(doc)
    add_heading(doc, "10. 보고서 본문 삽입용 종합 문안", 1)
    add_para(
        doc,
        "아래 문안은 특허 정제 방법론 또는 유효특허 선별 기준 항목에 그대로 사용하거나 분량에 맞춰 축약할 수 있다.",
    )

    add_heading(doc, "10.1 상세형", 2)
    detailed_text = (
        "본 분석에서는 휴머노이드 로봇 상용화와의 기술적 관련성을 기준으로 유효특허를 선별하였다. 먼저 동일 발명의 "
        "계속출원 및 패밀리 중복을 제거하고 출원인 표기를 정규화한 후, 제목·초록·대표청구항 및 AI 요약을 활용하여 "
        "후보군을 구성하였다. 이후 출원인이나 단순 키워드가 아니라 발명의 청구 대상과 핵심 기능을 기준으로 특허 단위 "
        "판정을 수행하였다. 휴머노이드·인간형·이족보행·전신제어를 직접 대상으로 하거나 완제품 휴머노이드의 핵심 "
        "서브시스템에 해당하는 특허는 Tier 1로 분류하였다. 휴머노이드 전용은 아니더라도 KIMM 가치사슬의 H/W·지능·데이터 "
        "축에 속하면서 로봇손·파지·촉각·힘제어, 균형·보행, 경량 액추에이터, VLA·모방학습, 텔레오퍼레이션·Sim-to-Real, "
        "인간협업 안전 등 휴머노이드로 이전 가능한 코어 기술은 Tier 2로 분류하였다. 반면 청구항이 용접·도장·웨이퍼·"
        "컨베이어·팔레타이징 등 특정 공정이나 설비에 묶이거나, 수술·외골격·청소·AGV·완구 등 특정 제품 목적에만 종속된 "
        "경우는 제외하였다. 다만 해당 분야의 특허라도 발명의 본체가 범용 QDD 관절, 임피던스 제어, 시연 데이터 수집 등 "
        "휴머노이드로 이전 가능한 코어 기술이면 Tier 2로 구제하였다. 최종적으로 Tier 1과 Tier 2를 분석상 유효특허로 "
        "포함하고, 특정 용도 전용 또는 키워드 오매칭 특허는 제외하였다."
    )
    p = add_para(doc, detailed_text, after=10, line=1.2)
    set_paragraph_shading(p, PALE_BLUE, BLUE)
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.08)

    add_heading(doc, "10.2 요약형", 2)
    short_text = (
        "유효특허는 출원인·키워드가 아니라 발명의 청구 대상과 핵심 기능을 기준으로 선별하였다. 휴머노이드·이족보행·"
        "전신제어 및 완제품 서브시스템은 Tier 1, H/W·지능·데이터 가치사슬에서 휴머노이드로 이전 가능한 조작·구동·학습·"
        "텔레오퍼레이션·안전 코어 기술은 Tier 2로 포함하였다. 특정 산업공정·설비·의료·외골격·청소·AGV·완구에 "
        "종속된 기술은 제외하되, 범용 코어 기술이 청구된 경우에는 Tier 2로 구제하였다. 동일 특허패밀리는 대표 1건으로 "
        "통합했으며, 최종 분석집합은 Tier 1과 Tier 2의 합으로 구성하였다."
    )
    p = add_para(doc, short_text, after=10, line=1.2)
    set_paragraph_shading(p, PALE_GREEN, GREEN)
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.08)

    add_heading(doc, "10.3 한 문장 요약", 2)
    add_callout(
        doc,
        "선별 기준",
        "휴머노이드에 직접 해당하거나 휴머노이드 상용화의 H/W·지능·데이터 병목에 이전 가능한 코어 기술은 포함하고, "
        "특정 공정·제품·응용례 또는 키워드에만 관련된 특허는 제외하였다.",
        fill=PALE_GOLD,
        accent=GOLD,
    )

    add_heading(doc, "11. 해석상의 유의사항", 1)
    add_bullet(
        doc,
        "사례 매핑은 카테고리별 후보 표본을 읽고 확정·경계·오탐을 구분한 판정 선례이며, 사례표 자체가 전수 라벨은 아니다.",
        bullet_num_id,
    )
    add_bullet(
        doc,
        "화이트리스트는 검토 우선순위를 높이는 장치이지 자동 포함 규칙이 아니다.",
        bullet_num_id,
    )
    add_bullet(
        doc,
        "경계 특허는 출원인보다 대표청구항의 전제부·필수 한정요소와 발명의 기술적 효과를 우선 확인해야 한다.",
        bullet_num_id,
    )
    add_bullet(
        doc,
        "최종 정답셋 확정 전에는 Tier별·제외군별 층화표본 검증을 수행하여 판정 일관성을 점검한다.",
        bullet_num_id,
    )

    add_heading(doc, "12. 근거 문서", 1)
    add_bullet(doc, "A2_정답셋_사례매핑.md — 카테고리별 확정·경계·오탐 사례와 판정 규칙 v1", bullet_num_id)
    add_bullet(doc, "A2_정답셋_구축계획.md — 데이터 범위, KIMM 3축, T1/T2/E 설계와 구축 절차", bullet_num_id)
    add_bullet(doc, "DataSet/humanoid/KIMM_핵심자료_정리.md — H/W·지능·데이터 가치사슬과 상용화 병목", bullet_num_id)
    add_bullet(doc, "DataSet/humanoid/A2_도메인설명.md — 도메인 정의, 핵심 기술 축, 포함·제외 의도", bullet_num_id)

    # Core properties and save.
    doc.core_properties.title = "휴머노이드 로봇 유효특허 선별 기준 및 판정 사례 보고서"
    doc.core_properties.subject = "A2 정답셋 구축을 위한 포함·제외 기준 v1"
    doc.core_properties.author = "A2 프로젝트팀"
    doc.core_properties.keywords = "휴머노이드, 특허, 유효특허, 선별기준, T1, T2, KIMM"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())

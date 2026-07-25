# -*- coding: utf-8 -*-
"""Create the final Gold v1.1 humanoid patent screening criteria report."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from create_patent_selection_report import (
    BLUE,
    DARK_BLUE,
    GOLD,
    GREEN,
    INK,
    LIGHT_BLUE,
    MUTED,
    NAVY,
    PALE_BLUE,
    PALE_GOLD,
    PALE_GREEN,
    PALE_RED,
    RED,
    add_bullet,
    add_callout,
    add_field,
    add_heading,
    add_kicker,
    add_label_line,
    add_numbering_definition,
    add_page_break,
    add_para,
    add_source_note,
    add_step,
    add_table,
    configure_styles,
    set_paragraph_shading,
    set_run_font,
)


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = (
    ROOT
    / "outputs"
    / "gold_v1_1_final"
    / "휴머노이드_Gold_v1_1_최종기준서.docx"
)


def configure_page_v11(doc):
    """Apply the standard_business_brief page geometry and v1.1 furniture."""
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    hr = hp.add_run("휴머노이드 특허 Gold v1.1  |  최종 선별 기준서")
    set_run_font(hr, size=8.5, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    fr = fp.add_run("gold v1.1  ·  2026.07.24  ·  ")
    set_run_font(fr, size=8.5, color=MUTED)
    page_run = fp.add_run()
    set_run_font(page_run, size=8.5, color=MUTED)
    add_field(page_run, "PAGE")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_metric_strip(doc):
    table = doc.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    labels = ["원자료", "패밀리 대표", "gold_label=1", "gold_label=0"]
    values = ["3,757건", "3,323건", "1,246건", "2,077건"]
    for col, label in enumerate(labels):
        cell = table.rows[0].cells[col]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        set_run_font(r, size=8.8, bold=True, color=NAVY)
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "D9EAF7")
        tc_pr.append(shd)
    for col, value in enumerate(values):
        cell = table.rows[1].cells[col]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(value)
        set_run_font(r, size=15, bold=True, color=NAVY if col < 2 else (GREEN if col == 2 else MUTED))
    widths = [2340, 2340, 2340, 2340]
    # Reuse the exact table geometry helper indirectly through add_table's module.
    from create_patent_selection_report import set_table_geometry, set_cell_borders, set_cell_margins

    set_table_geometry(table, widths)
    for row in table.rows:
        for cell in row.cells:
            set_cell_margins(cell, top=100, bottom=100, start=120, end=120)
            set_cell_borders(cell, color="B4C6E7", size=4)
    set_repeat_table_header(table.rows[0])
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    return table


def build_document():
    doc = Document()
    configure_styles(doc)
    configure_page_v11(doc)
    bullet_num_id = add_numbering_definition(doc, "bullet", "•", left=720, hanging=360)
    decimal_num_id = add_numbering_definition(doc, "decimal", "%1.", left=720, hanging=360)

    # memo_masthead opening, using the standard_business_brief preset.
    add_kicker(doc, "A2 정답셋 · FINAL TECHNICAL NOTE")
    title = add_para(
        doc,
        "휴머노이드 특허 Gold v1.1\n최종 선별 기준서",
        after=6,
        line=1.0,
        color=NAVY,
        keep=True,
    )
    for run in title.runs:
        set_run_font(run, size=24, bold=True, color=NAVY)
    subtitle = add_para(
        doc,
        "패밀리 대표 3,323건 · 원본 67열 복원 · 최종 유효특허 1,246건",
        after=14,
        line=1.1,
        color=MUTED,
        keep=True,
    )
    subtitle.runs[0].font.size = Pt(12.5)

    add_label_line(doc, "문서 목적  ", "보고서에 그대로 인용할 수 있는 유효특허 포함·제외 기준과 데이터 구축 절차의 확정본")
    add_label_line(doc, "판정 단위  ", "WIPS 패밀리 대표 1건")
    add_label_line(doc, "최종 산출물  ", "원본 67개 열 + gold_label + 기술분류 + 판정이유 = 70개 열")
    add_label_line(doc, "기준 상태  ", "gold v1.1 / 도메인 소유자 확정 규칙 12개 / 2026.07.24")
    add_metric_strip(doc)

    add_callout(
        doc,
        "최종 결론",
        "동일 패밀리를 대표 1건으로 통합한 3,323건 가운데, 휴머노이드에 직접 해당하거나 "
        "H/W·지능·데이터 가치사슬에서 휴머노이드로 이전 가능한 코어 기술 1,246건을 "
        "gold_label=1로 확정하였다. 특정 공정·제품·응용에 종속된 2,077건은 gold_label=0으로 제외하였다.",
        fill=PALE_GREEN,
        accent=GREEN,
    )
    add_callout(
        doc,
        "용어 주의",
        "이 문서의 ‘유효특허’는 본 연구의 기술 관련성 기준을 충족한 분석대상 특허를 뜻한다. "
        "특허권의 법률적 존속·무효 여부나 권리범위의 유효성을 판단한 것은 아니다.",
        fill=PALE_GOLD,
        accent=GOLD,
    )

    add_page_break(doc)
    add_heading(doc, "1. Gold v1.1의 구성과 확정 결과", 1)
    add_para(
        doc,
        "Gold v1.1은 원자료의 서지·초록·청구항·패밀리·인용·분류·AI 요약 등 67개 열을 그대로 보존하면서, "
        "분석용 라벨 3개 열을 뒤에 추가한 패밀리 대표 데이터셋이다. 원본 행을 임의 축약하거나 대표 필드를 "
        "재작성하지 않았으며, 출원번호를 주 키로 원자료와 결합하고 WIPS패밀리 ID로 교차검증하였다.",
    )
    add_table(
        doc,
        ["항목", "확정값", "의미"],
        [
            ("원자료", "3,757건 × 67열", "휴머노이드특허_Raw.xlsx의 전체 후보"),
            ("패밀리 대표", "3,323건", "동일 WIPS패밀리 중복 제거 후 분석 단위"),
            ("gold_label=1", "1,246건", "최종 유효특허"),
            ("gold_label=0", "2,077건", "최종 제외특허·대조군"),
            ("최종 스키마", "3,323건 × 70열", "원본 67열 + gold_label·기술분류·판정이유"),
        ],
        [1900, 1900, 5560],
        font_size=9.2,
        header_fill="D9EAF7",
    )
    add_source_note(
        doc,
        "검증: 최종 파일 재개방 후 행 3,323건, 열 70개, 양성 1,246건, 음성 2,077건, "
        "고유 패밀리 3,323개, 원본 67열 값 불일치 0건."
    )

    add_heading(doc, "1.1 v1에서 v1.1로의 최종 조정", 2)
    add_para(
        doc,
        "1차 판정의 양성 1,345건을 출발점으로 사례매핑과 확정 경계 규칙을 재적용하였다. "
        "산업용 주변기술 106건과 오탐 어휘 5건을 제외하고, 휴머노이드 OEM·직접 플랫폼 및 "
        "사례매핑 구제 대상 12건을 포함하여 1,246건으로 확정하였다.",
    )
    add_table(
        doc,
        ["조정", "건수", "내용", "누적 양성"],
        [
            ("초기 양성", "1,345", "1차 agentic 판정", "1,345"),
            ("제외", "-106", "일반제어·진동·경로·측정·교시UI·특정 공정 등 산업용 주변기술", "1,239"),
            ("제외", "-5", "밸런서·NC 등 오탐 사전 대상", "1,234"),
            ("구제", "+12", "사례매핑 T1/T2 및 OEM·직접 플랫폼 원칙", "1,246"),
        ],
        [1500, 1000, 5260, 1600],
        font_size=9.0,
    )

    add_page_break(doc)
    add_heading(doc, "2. 선별 절차와 의사결정 순서", 1)
    add_para(
        doc,
        "판정은 출원인 화이트리스트나 단순 키워드 점수로 자동 결정하지 않았다. 제목·초록·대표청구항·독립청구항·"
        "AI 요약을 함께 읽고, 청구 대상의 기술적 본체와 용도 전용성을 우선 판단하였다.",
    )
    for prefix, text in [
        ("1. ", "패밀리 정리 — WIPS패밀리 ID 기준으로 동일 발명을 대표 1건으로 통합한다."),
        ("2. ", "직접성 확인 — 휴머노이드·이족·인간형·전신제어·완제품 플랫폼을 직접 청구하면 포함한다."),
        ("3. ", "이전가능성 확인 — 휴머노이드 전용이 아니어도 H/W·지능·데이터 축의 코어 기술이면 포함한다."),
        ("4. ", "용도 전용성 확인 — 특정 공정·설비·제품 목적이 필수 한정요소이면 제외한다."),
        ("5. ", "경계 규칙 적용 — 4족, 물류 학습, 흡착, 외골격, 안전, 시뮬레이션 등 12개 확정 규칙을 적용한다."),
        ("6. ", "최종 기록 — gold_label, H1–H8 기술분류, 한국어 판정이유를 함께 남긴다."),
    ]:
        add_step(doc, text, decimal_num_id, bold_prefix=None, after=7)

    add_callout(
        doc,
        "핵심 판정식",
        "‘누가 출원했는가’보다 ‘무엇을 청구했는가’를 우선한다. 응용례에 휴머노이드가 등장하는 것만으로는 "
        "포함하지 않고, 반대로 산업용·수술·외골격 분야의 특허라도 범용 코어 기술 자체를 청구하면 구제한다.",
    )

    add_page_break(doc)
    add_heading(doc, "3. 기준서의 상위 뼈대: KIMM 가치사슬 3축", 1)
    add_para(
        doc,
        "최종 포함 범위는 KIMM 기계기술정책 No.122의 휴머노이드 가치사슬인 H/W·지능·데이터 3개 축에 "
        "앵커링하였다. H1–H8은 이 3축을 특허 판정에 사용할 수 있도록 기능 단위로 세분화한 체계다.",
    )
    add_table(
        doc,
        ["3축", "핵심 내용", "Gold v1.1에서의 대응"],
        [
            (
                "H/W",
                "액추에이터·감속기·제어시스템·센서·로봇손·배터리·열관리",
                "H2 로봇손, H3 제어, H4 균형·보행, H5 구동계, H8 안전 센싱·제어",
            ),
            (
                "지능",
                "VLA·E2E·RFM·SLAM·강화학습·AI칩·비전·제어 솔루션",
                "H3 학습·제어, H6 로봇비전·VLA·모방학습, H8 인간협업 판단",
            ),
            (
                "데이터",
                "시뮬레이터·Sim-to-Real·텔레오퍼레이션·시연/합성 데이터·미들웨어",
                "H7 텔레옵·Sim-to-Real·데이터",
            ),
        ],
        [1200, 3960, 4200],
        font_size=9.1,
        header_fill="D9EAF7",
    )

    add_page_break(doc)
    add_heading(doc, "4. 최종 기술분류 H1–H8", 1)
    add_para(
        doc,
        "gold_label=1인 특허에는 아래 H1–H8을 복수 부여한다. 하나의 특허가 로봇손 기구와 파지학습을 함께 "
        "청구하면 ‘H2; H3’처럼 기록한다. 따라서 분류별 건수 합계는 1,246건보다 클 수 있다.",
    )
    h_rows = [
        ("H1", "직접 휴머노이드", "통합", "휴머노이드·이족·인간형 본체, 전신제어, 완제품 플랫폼", "134"),
        ("H2", "로봇손·파지·촉각", "H/W", "다지 손, 그리퍼, 파지, 촉각·힘/토크 센싱, 인핸드 조작", "486"),
        ("H3", "매니퓰레이션 학습·제어", "H/W·지능", "모션 플래닝, 힘·임피던스 제어, 조작 스킬 학습", "623"),
        ("H4", "균형·보행·자세", "H/W", "ZMP, 전신동역학, 보행, 자세, 낙상 예측·복구", "415"),
        ("H5", "액추에이터·감속기·관절", "H/W", "QDD·SEA·VSA, 백드라이버블 관절, 감속기, 경량 구동계", "332"),
        ("H6", "로봇비전·VLA·모방학습", "지능", "로봇 비전, VLA, LfD, 강화학습 조작, 파지용 인식", "90"),
        ("H7", "텔레옵·Sim-to-Real·데이터", "데이터", "원격조작, 시연·합성 데이터, 도메인 무작위화, Sim-to-Real", "24"),
        ("H8", "인간협업 안전", "H/W·지능", "충돌·접촉 감지, 힘·속도 제한, 인간 공존 환경 안전", "171"),
    ]
    add_table(
        doc,
        ["코드", "분류", "3축", "포함 범위", "건수*"],
        h_rows,
        [800, 2200, 1300, 4060, 1000],
        font_size=8.7,
        header_fill="D9EAF7",
    )
    add_source_note(
        doc,
        "* 복수분류 기준. H1–H8 건수는 최종 엑셀의 기술분류 열을 기준으로 산출하며, "
        "gold_label=1인 1,246건 모두 최소 1개 이상의 H코드를 가진다."
    )

    add_heading(doc, "4.1 직접 휴머노이드와 이전가능 코어의 관계", 2)
    add_bullet(
        doc,
        "직접 휴머노이드: 휴머노이드·이족보행·인간형 전신, 본체·관절 구조, 전신제어, 낙상 대응 및 완제품 플랫폼을 포함한다.",
        bullet_num_id,
    )
    add_bullet(
        doc,
        "이전가능 코어: 휴머노이드 전용 표현이 없어도 H2–H8의 기능을 청구하며 특정 공정에 종속되지 않으면 포함한다.",
        bullet_num_id,
    )
    add_bullet(
        doc,
        "바퀴형 상반신 휴머노이드, 4족·다족·휠레그의 균형·구동 기술도 상용화 경로와 이족 이전성을 인정하여 포함할 수 있다.",
        bullet_num_id,
    )

    add_heading(doc, "5. 포함 기준", 1)
    inclusion_rows = [
        ("I1", "직접 폼팩터", "휴머노이드·인간형·이족·전신제어·완제품 플랫폼을 직접 청구"),
        ("I2", "조작·손", "다지 손, 파지·촉각·힘 센싱, 범용 그리퍼 및 인핸드 조작"),
        ("I3", "학습·제어", "다관절 조작 스킬, 모션 플래닝, 힘·임피던스·컴플라이언스 제어"),
        ("I4", "이동·균형", "전신 동역학, ZMP, 보행·자세, 외란·슬립·낙상 대응"),
        ("I5", "구동계", "QDD·SEA·VSA, 경량·고토크·백드라이버블 관절, 감속기"),
        ("I6", "지능", "로봇 비전, VLA, 모방학습, 강화학습 조작, 파지·자세 인식"),
        ("I7", "데이터", "텔레오퍼레이션, 시연 데이터 수집, 합성데이터, Sim-to-Real"),
        ("I8", "협업 안전", "인간 공존 환경의 충돌·접촉·속도·거리·힘 제한"),
    ]
    add_table(
        doc,
        ["코드", "포함 축", "판정 기준"],
        inclusion_rows,
        [900, 2100, 6360],
        font_size=9.0,
    )

    add_heading(doc, "6. 제외 기준", 1)
    exclusion_rows = [
        ("E1", "산업 공정·주변기술", "용접·도장·가공·웨이퍼·반도체 반송·SMT·교시UI·OLP·캘리브레이션 등"),
        ("E2", "수술·의료", "수술 셋업·카테터·의료 연속체 등 의료 목적에 청구 범위가 묶인 기술"),
        ("E3", "외골격·재활", "착용 보조·의지·의수·재활 자체가 목적이며 로봇 코어로 독립되지 않는 기술"),
        ("E4", "청소·잔디", "바닥청소·잔디깎기·충전도킹 등 청소 로봇 전용 기능"),
        ("E5", "물류·반송", "AGV·컨베이어·팔레타이징·분류 자원할당·물류 워크플로 전용"),
        ("E6", "완구·비실체 SW", "완구·애니매트로닉 또는 로봇 실체가 없는 가상화·통신 서비스"),
        ("E7", "오탐 어휘", "recycling, load balancing, gravity balancer, pedestrian, 가공물 pose 등 문맥 오인"),
    ]
    add_table(
        doc,
        ["코드", "제외 유형", "판정 기준"],
        exclusion_rows,
        [900, 2100, 6360],
        font_size=9.0,
        header_fill="F2F4F7",
    )
    add_callout(
        doc,
        "예외 원칙",
        "제외 분야에 속하더라도 발명의 청구 대상이 범용 QDD 관절, 임피던스 제어, 로봇 학습 데이터 수집, "
        "힘·촉각 기반 파지 등 독립적인 코어 기술이면 포함할 수 있다.",
        fill=PALE_BLUE,
        accent=BLUE,
    )

    add_page_break(doc)
    add_heading(doc, "7. Gold v1.1 확정 경계 규칙 12개", 1)
    rules = [
        (
            "1. 용도 전용성",
            "청구항 전제부·한정요소에 특정 공정·설비·용도가 등장하면 제외한다. 파지·힘제어·학습·관절 기구 자체를 청구하면 포함한다. 출원인이나 응용례 나열로 판단하지 않는다.",
        ),
        (
            "2. 완제품 OEM 3분할",
            "휴머노이드 제조사 특허도 제품 라인으로 나눈다. 휴머노이드 라인은 포함·핵심, legged/4족 일반은 포함·이전가능, 흡착 물류 전용 라인은 제외한다.",
        ),
        (
            "3. 역방향 구제",
            "자동차·전자 등 종합 대기업 특허라도 휴머노이드가 명시되면 핵심으로 포함한다.",
        ),
        (
            "4. 4족 원칙",
            "4족·다족·휠레그의 전신제어·다리 액추에이터·균형 기술은 이족 이전성을 인정하여 포함한다. 완구 4족만 제외한다.",
        ),
        (
            "5. 물류 학습 2분할",
            "빈피킹·파지 전략·접촉 기반 그립 조정처럼 파지·조작 스킬을 학습하면 포함한다. 팔레타이징 배치·분류 자원할당처럼 물류 워크플로를 학습하면 제외한다.",
        ),
        (
            "6. 산업용 회색지대",
            "힘 기반 직접교시·리드스루, 도메인 무작위화 성격의 시뮬레이션 파라미터 튜닝은 포함한다. OLP·공정 검증 시뮬레이터·교시 UI·캘리브레이션은 제외한다.",
        ),
        (
            "7. 흡착 3예외",
            "진공 흡착은 기본 제외하되 ①촉각/힘 피드백 실시간 조정 ②손가락 기구 하이브리드 ③파지품질 학습·물리모델이면 포함한다.",
        ),
        (
            "8. 소프트 그리퍼",
            "파지 원리·기구 특허는 포함(약)하고, 유체 배관·허브 등 시스템 패키징은 제외한다.",
        ),
        (
            "9. 외골격 예외",
            "착용형이라도 로봇 원격조작 마스터 또는 로봇 학습용 데이터 수집이 목적이면 포함한다. 가변강성 액추에이터 등 관절 요소기술 자체도 포함한다.",
        ),
        (
            "10. 안전 판별",
            "인간 공존 환경의 충돌·접촉·속도·거리 안전만 포함한다. 물체의 안전한 파지·이송 또는 펜스·인터록·고장 진단 같은 설비 기능안전은 해당 축에서 제외한다.",
        ),
        (
            "11. 시뮬레이션 판별",
            "학습·데이터 생성·Sim-to-Real·파라미터 튜닝용 시뮬레이션만 포함한다. 공정 검증·OLP·설계 도구는 제외한다.",
        ),
        (
            "12. 오탐 어휘 주의",
            "recycling은 재활이 아니고, load balancing은 균형 제어가 아니며, gravity balancer는 스프링 부품이다. pedestrian과 가공물 pose도 각각 인간협업·자세제어의 자동 근거가 아니다.",
        ),
    ]
    for label, text in rules:
        add_callout(doc, label, text, fill=PALE_BLUE if label[0] not in {"7", "9"} else PALE_GOLD, accent=BLUE if label[0] not in {"7", "9"} else GOLD)

    add_page_break(doc)
    add_heading(doc, "8. 대표 판정 사례", 1)
    add_heading(doc, "8.1 포함 사례", 2)
    add_table(
        doc,
        ["등록번호", "출원인", "기술 내용", "분류·판정"],
        [
            ("US12605824", "Figure AI", "휴머노이드 전신·엄지 직접구동 설계", "H1·H5 / 직접 핵심 포함"),
            ("US12172297", "Sanctuary", "힘줄·실리콘 스킨 기반 사람 손 모방 기계식 손", "H1·H2 / 포함"),
            ("US12318935", "NVIDIA", "신경 가치함수 기반 이동 물체 실시간 파지", "H3·H6 / 포함"),
            ("US11185975", "현대차·기아", "이족 이중지지 구간 구동력 제어", "H1·H4 / 포함"),
            ("US11865714", "삼성전자", "QDD 이중감속 준직접구동 관절", "H5 / 의료 오매칭 구제"),
            ("US12214495", "NAVER", "햅틱 원격 교시와 접촉 태스크 학습", "H3·H7 / 포함"),
            ("US11254004", "Veo Robotics", "정지 시간·거리 자동 추정", "H8 / 인간협업 안전 포함"),
        ],
        [1500, 1800, 4060, 2000],
        font_size=8.6,
        header_fill="D9EAF7",
    )

    add_heading(doc, "8.2 제외 사례", 2)
    add_table(
        doc,
        ["등록번호", "출원인", "기술 내용", "제외 사유"],
        [
            ("US9902070", "FANUC", "공작기계 냉각제 노즐 조정 로봇", "E1 / 공작기계 공정 전용"),
            ("US11285619", "FANUC", "칩마운터 테이프릴 파지 핸드", "E1 / SMT 설비 전용"),
            ("US11554498", "Kawasaki", "웨이퍼 지그·반송 기술", "E1 / 반도체 설비 전용"),
            ("US10603788", "FANUC", "오프라인 공정 검증 시뮬레이터", "E1 / 학습·Sim-to-Real 아님"),
            ("US12194634", "해군군의대·Surgerii", "수술 마스터-슬레이브 오류 검출", "E2 / 수술 전용"),
            ("US11833688", "PETOI", "장난감 4족 로봇", "E6 / 완구 4족"),
        ],
        [1500, 1800, 4060, 2000],
        font_size=8.6,
    )

    add_heading(doc, "8.3 같은 키워드가 갈리는 대비 사례", 2)
    add_table(
        doc,
        ["쟁점", "포함", "제외", "판정 포인트"],
        [
            (
                "산업용 로봇",
                "FANUC US11701777: 빈피킹 적응형 파지 계획",
                "FANUC US9902070: 공작기계 노즐 조정",
                "범용 파지 코어인가, 특정 공정 설비인가",
            ),
            (
                "시뮬레이션",
                "X Development US11458630: 접촉 시뮬레이션 기반 Sim-to-Real",
                "FANUC US10603788: OLP·공정 검증",
                "학습·데이터 생성인가, 공정 검증 도구인가",
            ),
            (
                "외골격",
                "Acumino US11822710: 로봇 학습용 인간 시연 데이터",
                "착용 보조·재활 자체 목적 특허",
                "착용 장치의 목적이 로봇 학습/텔레옵인가",
            ),
            (
                "안전",
                "Veo US11254004: 인간-로봇 거리·정지 안전",
                "Siemens US9707681: OLP 간섭 관리",
                "인간 공존 안전인가, 설비 기능안전인가",
            ),
        ],
        [1300, 2700, 2700, 2660],
        font_size=8.5,
        header_fill="D9EAF7",
    )
    add_source_note(
        doc,
        "대표 사례는 A2_정답셋_사례매핑.md의 확정·경계·오탐 판정 선례를 요약한 것이다. "
        "개별 특허의 최종 라벨은 Gold v1.1 데이터셋을 우선한다."
    )

    add_page_break(doc)
    add_heading(doc, "9. 최종 데이터 열의 해석", 1)
    add_table(
        doc,
        ["열", "허용값", "해석", "분석 시 주의"],
        [
            ("gold_label", "1", "최종 유효특허", "특허 통계·출원인 분석·클러스터링의 기본 모집단"),
            ("gold_label", "0", "최종 제외특허", "경계·오탐 검증용 대조군"),
            ("기술분류", "H1–H8", "포함 특허의 복수 기술축", "세미콜론 기준으로 분리 집계"),
            ("기술분류", "비대상(E0–E7)", "제외특허의 주된 제외 유형", "H1–H8 분포 집계에서 제외"),
            ("판정이유", "한국어 문장", "포함·제외 근거의 요약", "수동 조정은 ‘수동 구제/조정’으로 표시"),
        ],
        [1400, 1500, 3000, 3460],
        font_size=9.0,
        header_fill="D9EAF7",
    )
    add_para(
        doc,
        "기술분류는 판정 감사기록의 C1–C7 신호와 제목·초록·청구항·AI 요약의 강한 기능 신호를 결합해 부여하였다. "
        "포함 특허는 모두 최소 1개의 H코드를 갖고, 제외 특허는 H코드 대신 비대상(E코드)으로 구분하였다. "
        "gold_label 자체는 분류코드의 수나 빈도와 무관하게 v1.1 최종 판정을 그대로 사용한다.",
    )

    add_heading(doc, "10. 보고서 삽입용 문안", 1)
    add_heading(doc, "10.1 본문형", 2)
    report_text = (
        "본 연구는 휴머노이드 로봇 관련 미국 특허 후보 3,757건을 대상으로 동일 WIPS 특허패밀리의 중복을 제거하여 "
        "3,323건의 패밀리 대표 특허를 분석 단위로 구성하였다. 유효특허 선별은 출원인 또는 단순 키워드가 아니라 "
        "제목·초록·대표 및 독립청구항, AI 요약에 나타난 발명의 청구 대상과 기술적 핵심을 기준으로 수행하였다. "
        "휴머노이드·인간형·이족보행·전신제어·완제품 플랫폼을 직접 대상으로 하는 특허와, 휴머노이드 전용은 아니더라도 "
        "KIMM 가치사슬의 H/W·지능·데이터 축에서 로봇손·파지·촉각, 매니퓰레이션 학습·제어, 균형·보행, "
        "액추에이터·감속기·관절, 로봇비전·VLA·모방학습, 텔레오퍼레이션·Sim-to-Real, 인간협업 안전에 해당하는 "
        "이전가능 코어 기술을 포함하였다. 반면 청구항이 용접·도장·가공·웨이퍼·컨베이어·팔레타이징 등 특정 산업 "
        "공정이나 설비에 묶이거나 수술·외골격·청소·물류 반송·완구·비실체 소프트웨어 등 특정 제품 목적에 종속된 "
        "경우는 제외하였다. 다만 제외 분야의 특허라도 범용 QDD 관절, 힘·임피던스 제어, 로봇 학습 데이터 수집, "
        "힘·촉각 기반 파지처럼 휴머노이드로 이전 가능한 코어 기술 자체를 청구한 경우에는 포함하였다. 이 기준에 따라 "
        "3,323건 중 1,246건을 최종 유효특허(gold_label=1), 2,077건을 제외특허(gold_label=0)로 확정하였다."
    )
    p = add_para(doc, report_text, after=10, line=1.2)
    set_paragraph_shading(p, PALE_BLUE, BLUE)
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.08)

    add_heading(doc, "10.2 요약형", 2)
    short_text = (
        "동일 특허패밀리를 대표 1건으로 통합한 3,323건을 대상으로, 청구 대상과 기술적 핵심이 휴머노이드에 직접 "
        "해당하거나 H/W·지능·데이터 가치사슬에서 휴머노이드로 이전 가능한 코어 기술인 경우를 포함하였다. 특정 "
        "공정·설비·의료·외골격·청소·물류·완구에 종속된 기술은 제외하되, 범용 코어 기술 자체가 청구된 경우에는 "
        "예외적으로 구제하였다. 최종 유효특허는 1,246건이다."
    )
    p = add_para(doc, short_text, after=10, line=1.2)
    set_paragraph_shading(p, PALE_GREEN, GREEN)
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.08)

    add_heading(doc, "10.3 한 문장형", 2)
    add_callout(
        doc,
        "선별 기준",
        "휴머노이드에 직접 해당하거나 H/W·지능·데이터 상용화 병목에 이전 가능한 코어 기술은 포함하고, "
        "특정 공정·제품·응용 또는 키워드에만 관련된 특허는 제외하였다.",
        fill=PALE_GOLD,
        accent=GOLD,
    )

    add_heading(doc, "11. 사용 범위와 한계", 1)
    add_bullet(
        doc,
        "분석 모집단은 현재 확보한 미국 특허 원자료와 검색 범위에 한정되므로, 전 세계 휴머노이드 특허의 완전한 전수집합을 의미하지 않는다.",
        bullet_num_id,
    )
    add_bullet(
        doc,
        "패밀리 대표 1건 기준이므로 문헌 수를 세는 분석과 발명 수를 세는 분석을 구분해야 한다. 본 Gold v1.1은 후자에 적합하다.",
        bullet_num_id,
    )
    add_bullet(
        doc,
        "H1–H8은 복수분류이며 분류 합계가 1,246건을 초과한다. 단일 클러스터가 필요하면 연구 목적에 맞는 우선순위 규칙을 별도로 적용한다.",
        bullet_num_id,
    )
    add_bullet(
        doc,
        "gold_label은 확정 라벨이고, 판정이유는 감사 편의를 위한 압축 설명이다. 법률적 권리 유효성·침해·무효 판단에 사용하지 않는다.",
        bullet_num_id,
    )
    add_bullet(
        doc,
        "신규 특허를 추가할 때는 동일한 패밀리 대표 원칙과 12개 경계 규칙을 적용하고, 기존 v1.1 행의 라벨을 임의 변경하지 않는다.",
        bullet_num_id,
    )

    add_heading(doc, "12. 근거 문서와 데이터 계보", 1)
    add_table(
        doc,
        ["구분", "파일", "역할"],
        [
            ("원자료", "휴머노이드문제/휴머노이드특허_Raw.xlsx", "원본 3,757건·67열"),
            ("최종 라벨", "DataSet/humanoid/goldset_v1_1.csv", "패밀리 대표 3,323건·최종 gold v1.1"),
            ("판정 감사", "DataSet/.../judge/audit.jsonl", "C1–C7·E 판정 및 rationale"),
            ("확정 규칙", "DataSet/humanoid/A2_판정규칙_v1.md", "도메인 소유자 확정 12개 규칙"),
            ("사례 매핑", "A2_정답셋_사례매핑.md", "확정·경계·오탐 판정 사례"),
            ("KIMM 근거", "DataSet/humanoid/KIMM_핵심자료_정리.md", "H/W·지능·데이터 3축"),
            ("최종 엑셀", "휴머노이드_Gold_v1_1_패밀리대표_3323건.xlsx", "원본 67열 + 분석 열 3개"),
        ],
        [1400, 3860, 4100],
        font_size=8.8,
        header_fill="D9EAF7",
    )
    add_source_note(
        doc,
        "기준 간 충돌 시 우선순위: A2_판정규칙_v1.md(도메인 소유자 확정) → goldset_v1_1.csv(최종 라벨) "
        "→ 사례매핑·판정 감사 기록 → 자동 분류 신호."
    )

    doc.core_properties.title = "휴머노이드 특허 Gold v1.1 최종 선별 기준서"
    doc.core_properties.subject = "패밀리 대표 3,323건 및 최종 유효특허 1,246건의 포함·제외 기준"
    doc.core_properties.author = "A2 프로젝트팀"
    doc.core_properties.keywords = "휴머노이드, 특허, Gold v1.1, 유효특허, H1-H8, KIMM, 패밀리"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())

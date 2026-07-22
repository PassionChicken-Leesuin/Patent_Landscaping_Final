from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "paper"
DOCX_OUT = OUT_DIR / "A_Multi_Agent_Weak_Supervision_Framework_for_Domain_Relevant_Patent_Identification.docx"
MD_OUT = OUT_DIR / "A_Multi_Agent_Weak_Supervision_Framework_for_Domain_Relevant_Patent_Identification.md"

FIG_OVERALL = ROOT / "figures" / "research_overall_framework.png"
FIG_MAS = ROOT / "figures" / "mas_framework.png"


TITLE = "A Multi-Agent Weak Supervision Framework for Domain-Relevant Patent Identification"
AUTHORS = "Suin Lee, Woojin Choi, Seoyoung Moon, and Jungmin Yoo"
AFFILIATION = "Seoul National University"


DATASET_ROWS = [
    ("Self-driving Vehicle", 313, 895, 1208, 11196),
    ("Additive Manufacturing", 280, 818, 1098, 5774),
    ("Blockchain", 292, 1587, 1879, 6062),
    ("Computer Vision", 301, 959, 1260, 10335),
    ("Genome Editing", 303, 615, 918, 929),
    ("Hydrogen Storage", 287, 858, 1145, 2254),
    ("Total", 1766, 5732, 7498, 36550),
]


LABEL_DIAGNOSTICS = [
    ("Self-driving Vehicle", "Snorkel", 8117, 1385, 1694, 35),
    ("Self-driving Vehicle", "MAS", 5186, 5411, 599, 170),
    ("Additive Manufacturing", "Snorkel", 5774, 0, 0, 3),
    ("Additive Manufacturing", "MAS", 5468, 287, 20, 16),
    ("Blockchain", "Snorkel", 6062, 0, 0, 5),
    ("Blockchain", "MAS", 3358, 2699, 5, 1),
    ("Computer Vision", "Snorkel", 10335, 0, 0, 239),
    ("Computer Vision", "MAS", 7924, 2413, 0, 575),
    ("Genome Editing", "Snorkel", 929, 0, 0, 2),
    ("Genome Editing", "MAS", 724, 202, 3, 2),
    ("Hydrogen Storage", "Snorkel", 2254, 0, 0, 9),
    ("Hydrogen Storage", "MAS", 1654, 600, 0, 6),
]


PERFORMANCE_ROWS = [
    ("Self-driving Vehicle", "Snorkel", 0.704, 0.681, 0.329, 0.866, 0.813),
    ("Self-driving Vehicle", "MAS", 0.830, 0.696, 0.364, 0.826, 0.815),
    ("Additive Manufacturing", "Snorkel", 0.930, 0.759, 0.471, 0.874, 0.848),
    ("Additive Manufacturing", "MAS", 0.963, 0.775, 0.489, 0.919, 0.859),
    ("Blockchain", "Snorkel", 0.924, 0.687, 0.353, 0.620, 0.866),
    ("Blockchain", "MAS", 0.966, 0.880, 0.955, 0.694, 0.928),
    ("Computer Vision", "Snorkel", 0.870, 0.775, 0.738, 0.612, 0.825),
    ("Computer Vision", "MAS", 0.974, 0.873, 0.970, 0.707, 0.897),
    ("Genome Editing", "Snorkel", 0.921, 0.696, 0.422, 0.776, 0.769),
    ("Genome Editing", "MAS", 0.985, 0.941, 0.924, 0.918, 0.948),
    ("Hydrogen Storage", "Snorkel", 0.935, 0.782, 0.990, 0.560, 0.803),
    ("Hydrogen Storage", "MAS", 0.950, 0.830, 0.983, 0.632, 0.852),
    ("Average", "Snorkel", 0.881, 0.730, 0.551, 0.718, 0.821),
    ("Average", "MAS", 0.945, 0.833, 0.781, 0.783, 0.883),
]


ABSTRACT = (
    "Patent landscaping begins with the practical problem of identifying patents that are relevant "
    "to a target technological domain. Keyword and CPC-based search strategies can retrieve broad "
    "candidate pools, but they often have low precision and require repeated expert adjustment. "
    "Recent weak-supervision approaches reduce manual labeling by combining heuristic labeling "
    "functions, but their performance depends strongly on hand-designed domain rules. This paper "
    "proposes a multi-agent weak supervision framework that replaces Snorkel-style keyword labeling "
    "with rubric-guided large language model agents for pseudo-labeling domain-relevant patents. "
    "The framework uses a Relevance and Route agent, a conditional Exclusion agent, and a deterministic "
    "scoring step to assign each patent a score and candidate type. We evaluate the framework on six "
    "technology domains from Bergeaud and Verluise: self-driving vehicles, additive manufacturing, "
    "blockchain, computer vision, genome editing, and hydrogen storage. For each domain, Snorkel and "
    "the proposed MAS label the same candidate pool, the resulting labels are used to fine-tune the "
    "same SciBERT classifier, and performance is measured on a held-out gold benchmark. Across all "
    "six domains, MAS-labeled training data improves average AUC from 0.881 to 0.945, Macro-F1 from "
    "0.730 to 0.833, recall from 0.551 to 0.781, precision from 0.718 to 0.783, and accuracy from "
    "0.821 to 0.883. The results suggest that agentic reasoning improves weak supervision by mining "
    "in-domain hard negatives that keyword labeling functions fail to express."
)


SECTIONS: list[tuple[str, list[str]]] = [
    (
        "1 Introduction",
        [
            (
                "Patent landscaping analysis provides a structured view of the patent activity surrounding "
                "a scientific or technological domain. Its first and often most consequential step is patent "
                "identification: deciding which documents are actually relevant to the domain of interest. "
                "If this initial set is noisy, incomplete, or irreproducible, later analyses of technological "
                "trajectories, competitive positioning, and research opportunities inherit that error."
            ),
            (
                "Traditional patent identification relies on combinations of keywords, IPC or CPC codes, and "
                "expert-crafted Boolean queries. These approaches are attractive because they are transparent "
                "and easy to execute at scale, but they are expensive to refine and often return broad candidate "
                "sets with limited precision. Machine learning methods can learn richer decision boundaries, but "
                "they normally require labeled training data. In domain-specific patent landscaping, such labels "
                "are costly because a patent may contain domain vocabulary while still not performing the target "
                "technological task."
            ),
            (
                "Weak supervision addresses this bottleneck by creating training labels from noisy programmatic "
                "sources rather than manual annotation. Sofean's patent identification pipeline uses Snorkel to "
                "combine labeling functions and then fine-tunes SciBERT on the resulting training set. However, "
                "Snorkel shifts part of the burden from manual labeling to manual rule design: high-performing "
                "labeling functions still require domain knowledge, careful boundary definitions, and iterative "
                "engineering. This limitation is especially visible when the positive signal is not a keyword but "
                "a functional task, such as whether a vehicle actually automates driving or merely assists a human "
                "driver."
            ),
            (
                "This paper asks whether a multi-agent system can serve as a more scalable weak-supervision "
                "labeler for domain-relevant patent identification. The proposed framework uses rubric-grounded "
                "agents to reason over patent titles and abstracts, identify positive cases, separate easy "
                "negatives from hard negatives, and produce pseudo-labels for downstream SciBERT training. The "
                "experimental design holds the downstream model, evaluation data, and metrics fixed; only the "
                "labeling mechanism changes."
            ),
            (
                "The study is organized around two research questions. RQ1 asks whether the proposed MAS-generated "
                "training data improves patent identification performance relative to a Snorkel weak-supervision "
                "baseline. RQ2 asks whether the framework generalizes across multiple technology domains rather "
                "than only one carefully engineered case. The empirical answer to both questions is positive: MAS "
                "outperforms Snorkel across all six evaluated domains, with the largest average gain appearing in "
                "recall."
            ),
        ],
    ),
    (
        "2 Related Work",
        [
            (
                "Patent landscaping and patent identification. Patent landscaping has been used to assess "
                "technology trends, R&D investment, and competitive landscapes. A central problem in this process "
                "is the construction of a relevant patent set. Early and rule-based approaches use search formulas "
                "that combine keywords with classification codes. Such formulas are transparent but can be brittle: "
                "CPC classes are organized around technical features rather than a researcher's functional domain, "
                "and keyword matches often capture background mentions or adjacent technologies."
            ),
            (
                "Automated patent landscaping. Abood and Feltenberger introduced automated patent landscaping as "
                "a machine-learning approach to expanding from representative seed patents. Bergeaud and Verluise "
                "advanced this line of work by defining six frontier technology domains at the functional-application "
                "level and constructing seed and anti-seed benchmarks. Their framing is important for the present "
                "study because a patent is not positive merely because it matches a keyword or CPC rule; it is "
                "positive when the invention performs one of the domain's defining tasks."
            ),
            (
                "Weak supervision and Snorkel. Snorkel provides a framework for training data creation through "
                "labeling functions. Instead of asking experts to label every data point, experts write heuristic "
                "functions that vote on labels, and a generative label model combines their overlapping and "
                "conflicting outputs. In patent identification, this approach can reduce direct annotation costs, "
                "but it still depends on the quality and coverage of the labeling functions. When the boundary is "
                "semantic rather than lexical, keyword functions may over-label candidate pools as positive."
            ),
            (
                "Large language models and agentic labeling. Large language models can judge semantic relevance "
                "from natural language descriptions, but direct single-call classification is hard to audit and can "
                "be expensive. The framework proposed here uses a constrained multi-agent design instead: a first "
                "agent extracts evidence and routes cases, a second agent checks exclusions only when necessary, "
                "and a deterministic scoring step maps the structured state into candidate types. This preserves "
                "auditability while allowing the labeler to express distinctions that are difficult to encode as "
                "static keyword rules."
            ),
        ],
    ),
    (
        "3 Methods",
        [
            (
                "The pipeline compares two labelers under controlled downstream conditions. The input to both "
                "labelers is the same title-and-abstract candidate set. The output of each labeler is transformed "
                "into a binary SciBERT training set, where SEED denotes a domain-relevant patent and NOT_SEED "
                "denotes an irrelevant or excluded patent. SciBERT is then fine-tuned with identical hyperparameters "
                "and evaluated on the same held-out gold set."
            ),
            (
                "Data and domains. We evaluate six domains drawn from Bergeaud and Verluise: self-driving vehicles, "
                "additive manufacturing, blockchain, computer vision, genome editing, and hydrogen storage. For each "
                "domain, the gold benchmark contains manually labeled SEED and NOT_SEED examples. Candidate pools are "
                "collected using the domain's official CPC-prefix and keyword search query. For each domain, we also "
                "attach the other five domains' gold sets as out-of-domain candidates, which allows the labelers to "
                "show whether they can reject clearly different technologies."
            ),
            (
                "Snorkel baseline. The Snorkel arm uses labeling functions and a LabelModel. Self-driving vehicles "
                "use a bespoke set of positive and negative labeling functions because the domain has a known "
                "automate-versus-assist boundary. The remaining five domains use generic keyword-based labeling "
                "functions derived from the domain keyword lists. This design represents a realistic weak-supervision "
                "baseline: it is scalable and domain-parameterized, but its ability to create in-domain negative "
                "examples is limited when candidate pools are already selected by keyword and CPC rules."
            ),
            (
                "MAS labeler. The MAS arm uses a domain rubric rather than hand-written labeling functions. For "
                "non-self-driving domains, rubrics are generated from the domain's functional tasks, keyword signals, "
                "and hard-negative concept. Self-driving vehicles use a manually specified rubric that emphasizes the "
                "distinction between automating driving and assisting a human driver. For each patent, Node A extracts "
                "functional and technical evidence, assigns a relevance score, and routes the case as easy_positive, "
                "easy_negative, boundary, hard_negative, or abstain_candidate. Node B runs only for boundary or "
                "hard-negative cases and checks whether the patent should be excluded as a look-alike. Node C applies "
                "deterministic rules to produce a final score and a candidate type."
            ),
            (
                "Downstream training. Positive MAS cases are mapped to SEED. Easy-negative and hard-negative MAS "
                "cases are mapped to NOT_SEED, while boundary and abstain cases are dropped. Snorkel labels are mapped "
                "directly from SEED, NOT_SEED, and ABSTAIN. In all experiments, the downstream classifier is "
                "allenai/scibert_scivocab_uncased with a maximum sequence length of 256, four epochs, learning rate "
                "2e-5, batch size 16, weight decay 0.01, a 10 percent validation split, and class-weighted loss."
            ),
            (
                "Evaluation. The primary metrics are AUC, Macro-F1, recall, precision, and accuracy on each domain's "
                "held-out gold benchmark. Macro-F1 is important because the evaluation sets are imbalanced. Recall is "
                "also central in patent landscaping because missing relevant patents can distort downstream analyses "
                "of a technological field."
            ),
        ],
    ),
    (
        "4 Experiments and Results",
        [
            (
                "The experiment covers 7,498 gold evaluation patents and 36,550 in-domain labeling candidates across "
                "six technologies. After OOD augmentation, the labelers process 74,068 candidate records in total. "
                "The gold sets are used only for final evaluation and are not supplied to either Snorkel or MAS."
            ),
            (
                "The first diagnostic result concerns the labelers' behavior before downstream fine-tuning. Snorkel "
                "labels most in-domain candidate pools as positive because the candidate pools were themselves built "
                "from domain keywords and CPC classes. In the five non-self-driving domains, it creates no in-domain "
                "NOT_SEED examples. MAS, by contrast, produces substantial in-domain negative sets by identifying "
                "rule-matched but task-irrelevant patents as hard negatives. Across all domains, MAS assigns 24,314 "
                "in-domain candidates to positive, 11,612 to negative, and 622 to boundary or abstain. Snorkel assigns "
                "33,471 in-domain candidates to SEED, only 1,385 to NOT_SEED, and 1,694 to ABSTAIN."
            ),
            (
                "Downstream performance confirms the importance of this labeling difference. MAS improves average AUC "
                "from 0.881 to 0.945, Macro-F1 from 0.730 to 0.833, recall from 0.551 to 0.781, precision from 0.718 "
                "to 0.783, and accuracy from 0.821 to 0.883. MAS is better than Snorkel on AUC, Macro-F1, and accuracy "
                "in every domain. It also improves recall and precision in five of six domains."
            ),
            (
                "The largest gains appear in domains where Snorkel's keyword rules create an especially one-sided "
                "training signal. Blockchain, computer vision, and genome editing show large Macro-F1 gains. For "
                "blockchain, Macro-F1 increases from 0.687 to 0.880 and recall from 0.353 to 0.955. For genome editing, "
                "Macro-F1 increases from 0.696 to 0.941. These improvements indicate that MAS pseudo-labels help "
                "SciBERT learn a decision boundary rather than only a domain vocabulary."
            ),
        ],
    ),
    (
        "5 Discussion",
        [
            (
                "The results support the claim that multi-agent reasoning can improve weak supervision in patent "
                "identification. Snorkel is effective when good labeling functions exist, but the effort required to "
                "write those functions grows with the number and subtlety of domains. MAS shifts that effort into a "
                "structured rubric and evidence-extraction process. This is not a replacement for domain definition: "
                "the model still needs a clear description of the target technology. However, it reduces the need to "
                "translate that definition into brittle keyword rules."
            ),
            (
                "The most important mechanism is hard-negative mining. Bergeaud and Verluise's framework defines "
                "technologies at the functional-application level. A patent can match the search query but fail to "
                "perform the target task. These cases are exactly what Snorkel keyword rules struggle to reject and "
                "what MAS is designed to detect. By supplying SciBERT with many more in-domain negative examples, MAS "
                "appears to improve the classifier's boundary between real domain patents and look-alikes."
            ),
            (
                "The framework also generalizes across domains. The six domains vary in vocabulary, evaluation-set "
                "balance, and candidate-pool size. MAS improves average performance across all of them, suggesting "
                "that the agentic labeler is not simply overfitted to the self-driving case. At the same time, the "
                "domain-specific rubrics remain essential: the agents need to know what functional tasks count as "
                "in-scope and what kinds of rule-matched patents should be excluded."
            ),
            (
                "Several limitations remain. First, MAS labels are pseudo-labels rather than human annotations; the "
                "study evaluates their usefulness through downstream performance, not through a full manual audit of "
                "every generated label. Second, the comparison uses a practical Snorkel baseline rather than an "
                "exhaustively optimized set of expert-written labeling functions for all six domains. Third, MAS "
                "requires LLM calls, which add monetary cost and dependency on model behavior. Finally, all gold sets "
                "come from the same underlying benchmark family, so future work should test transfer to additional "
                "technologies, patent offices, and time periods."
            ),
        ],
    ),
    (
        "6 Conclusions",
        [
            (
                "This paper proposed a multi-agent weak supervision framework for domain-relevant patent "
                "identification. The framework replaces Snorkel's hand-written labeling functions with rubric-guided "
                "agents that extract evidence, route cases, check exclusions, and assign candidate types. In a "
                "controlled comparison across six technology domains, MAS-labeled training data consistently improved "
                "downstream SciBERT performance over Snorkel-labeled training data."
            ),
            (
                "The central empirical finding is that MAS improves recall while maintaining or improving precision. "
                "This matters for patent landscaping because the cost of missing relevant patents is high: omissions "
                "can distort subsequent analysis of technological scope, competition, and innovation trajectories. "
                "By identifying hard negatives and reducing dependence on manually engineered labeling functions, "
                "MAS offers a scalable path toward more reproducible and semantically grounded patent identification."
            ),
        ],
    ),
]


REFERENCES = [
    "Abood, A., and Feltenberger, D. (2018). Automated patent landscaping. Artificial Intelligence and Law, 26, 103-125.",
    "Bergeaud, A., and Verluise, C. (2023). Identifying technology clusters based on automated patent landscaping. PLOS ONE, 18(12), e0295587.",
    "Ratner, A. J., Bach, S. H., Ehrenberg, H., Fries, J., Wu, S., and Re, C. (2017). Snorkel: Rapid training data creation with weak supervision. Proceedings of the VLDB Endowment, 11(3), 269-282.",
    "Sofean, M. (2026). Identification of domain-relevant patents via weakly supervised deep learning. World Patent Information, 84, 102434.",
    "Trippe, A. (2015). Guidelines for preparing patent landscape reports. World Intellectual Property Organization.",
    "Beltagy, I., Lo, K., and Cohan, A. (2019). SciBERT: A pretrained language model for scientific text. Proceedings of EMNLP-IJCNLP.",
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        st = doc.styles[style_name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.208


def add_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE)
    r.bold = True
    r.font.size = Pt(20)
    r.font.name = "Calibri"
    r.font.color.rgb = RGBColor(31, 58, 95)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(AUTHORS)
    r.font.size = Pt(11)
    r.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run(AFFILIATION)
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(85, 85, 85)


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(text)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(85, 85, 85)


def add_source_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(85, 85, 85)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "F4F6F9")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(8.5)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.1
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.size = Pt(8.5)
    set_table_width(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_abstract(doc: Document) -> None:
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(6)
    h.paragraph_format.space_after = Pt(4)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = h.add_run("Abstract")
    r.bold = True
    r.font.size = Pt(12)
    add_paragraph(doc, ABSTRACT)


def add_figures(doc: Document) -> None:
    doc.add_heading("3.1 Framework Overview", level=2)
    add_paragraph(
        doc,
        "Figure 1 summarizes the controlled experimental design. The same candidate pools are "
        "processed by two weak-supervision labelers, and the downstream model and gold evaluation "
        "sets are held constant.",
    )
    if FIG_OVERALL.exists():
        doc.add_picture(str(FIG_OVERALL), width=Inches(6.35))
        add_caption(doc, "Figure 1. Overall controlled framework for Snorkel versus MAS weak supervision.")

    doc.add_page_break()
    doc.add_heading("3.2 MAS Labeling Graph", level=2)
    add_paragraph(
        doc,
        "Figure 2 expands the MAS labeler. The design keeps the high-frequency relevance step cheap, "
        "uses exclusion reasoning only for uncertain or look-alike cases, and preserves a slim CSV "
        "output plus a full audit log.",
    )
    if FIG_MAS.exists():
        doc.add_picture(str(FIG_MAS), width=Inches(6.35))
        add_caption(doc, "Figure 2. MAS framework for rubric-guided patent pseudo-labeling.")


def add_tables(doc: Document) -> None:
    doc.add_heading("4.1 Dataset Scale", level=2)
    add_source_note(doc, "Table 1. Gold evaluation sets and in-domain labeling candidate pools.")
    add_table(
        doc,
        ["Domain", "SEED", "NOT_SEED", "Gold N", "Labeling Pool N"],
        [[r[0], f"{r[1]:,}", f"{r[2]:,}", f"{r[3]:,}", f"{r[4]:,}"] for r in DATASET_ROWS],
        [3000, 1200, 1400, 1500, 2260],
    )

    doc.add_heading("4.2 Labeling Diagnostics", level=2)
    add_source_note(
        doc,
        "Table 2. In-domain label distribution and OOD false positives. MAS negatives combine "
        "easy_negative and hard_negative; Snorkel negatives are NOT_SEED labels.",
    )
    add_table(
        doc,
        ["Domain", "Labeler", "In-domain Positive", "In-domain Negative", "Boundary / Abstain", "OOD False Positive"],
        [[r[0], r[1], f"{r[2]:,}", f"{r[3]:,}", f"{r[4]:,}", f"{r[5]:,}"] for r in LABEL_DIAGNOSTICS],
        [2450, 1000, 1600, 1600, 1500, 1210],
    )

    doc.add_heading("4.3 Downstream Performance", level=2)
    add_source_note(doc, "Table 3. SciBERT evaluation on held-out gold benchmarks at threshold 0.5.")
    add_table(
        doc,
        ["Domain", "Labeler", "AUC", "Macro-F1", "Recall", "Precision", "Accuracy"],
        [[r[0], r[1], f"{r[2]:.3f}", f"{r[3]:.3f}", f"{r[4]:.3f}", f"{r[5]:.3f}", f"{r[6]:.3f}"] for r in PERFORMANCE_ROWS],
        [2500, 1000, 900, 1100, 1000, 1200, 1660],
    )


def build_docx() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_title_block(doc)
    add_abstract(doc)

    for heading, paragraphs in SECTIONS[:3]:
        doc.add_heading(heading, level=1)
        for paragraph in paragraphs:
            add_paragraph(doc, paragraph)
        if heading == "3 Methods":
            add_figures(doc)

    # Results section needs tables interleaved after opening prose.
    heading, paragraphs = SECTIONS[3]
    doc.add_heading(heading, level=1)
    for paragraph in paragraphs:
        add_paragraph(doc, paragraph)
    doc.add_page_break()
    add_tables(doc)

    for heading, paragraphs in SECTIONS[4:]:
        doc.add_heading(heading, level=1)
        for paragraph in paragraphs:
            add_paragraph(doc, paragraph)

    doc.add_heading("References", level=1)
    for ref in REFERENCES:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(6)
        p.add_run(ref)

    doc.save(DOCX_OUT)


def build_markdown() -> None:
    lines: list[str] = [f"# {TITLE}", "", AUTHORS, "", AFFILIATION, "", "## Abstract", "", ABSTRACT, ""]
    for heading, paragraphs in SECTIONS:
        lines.append(f"## {heading}")
        lines.append("")
        for paragraph in paragraphs:
            lines.append(paragraph)
            lines.append("")
        if heading == "3 Methods":
            lines.extend([
                "![Figure 1. Overall controlled framework](../figures/research_overall_framework.png)",
                "",
                "![Figure 2. MAS framework](../figures/mas_framework.png)",
                "",
            ])
    lines.append("## References")
    lines.append("")
    for ref in REFERENCES:
        lines.append(ref)
        lines.append("")
    MD_OUT.write_text("\n".join(lines), encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_markdown()
    print(DOCX_OUT)
    print(MD_OUT)

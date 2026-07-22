from __future__ import annotations

import shutil
import tempfile
import uuid
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from generate_paper_draft import (
    ABSTRACT,
    AFFILIATION,
    AUTHORS,
    DATASET_ROWS,
    LABEL_DIAGNOSTICS,
    PERFORMANCE_ROWS,
    REFERENCES,
    SECTIONS,
    TITLE,
)


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "paper" / "A_Multi_Agent_Weak_Supervision_Framework_for_Domain_Relevant_Patent_Identification_final.docx"
FIG_OVERALL = ROOT / "figures" / "research_overall_framework_bw.png"
FIG_MAS = ROOT / "figures" / "mas_framework_bw.png"
FONT_DIR = ROOT / "paper" / "fonts" / "cmu_ttf" / "cm-unicode-0.7.0"
FONT_NAME = "CMU Serif"


FINAL_REFERENCES = [
    "Abood, A., & Feltenberger, D. (2018). Automated patent landscaping. Artificial Intelligence and Law, 26(2), 103-125.",
    "Beltagy, I., Lo, K., & Cohan, A. (2019). SciBERT: A pretrained language model for scientific text. Proceedings of EMNLP-IJCNLP.",
    "Bergeaud, A., Potiron, Y., & Raimbault, J. (2017). Classifying patents based on their semantic content. PLOS ONE, 12(4), e0176310.",
    "Choi, S., Lee, H., Park, E. L., & Choi, S. (2019). Deep patent landscaping model using transformer and graph embedding. arXiv preprint arXiv:1903.05823.",
    "Clarke, N. S., Jürgens, B., & Herrero-Solana, V. (2020). Blockchain patent landscaping: An expert based methodology and search query. World Patent Information, 61, 101964.",
    "Cypris. (2026, March 2). How to do a patent landscape analysis in the age of AI. https://www.cypris.ai/insights/how-to-do-a-patent-landscape-analysis-in-the-age-of-ai",
    "Green, M., Halstead, M., Jay, C., Kingston, R., Singleton, A., & Topping, D. (2026). Comparing how large language models perform against keyword-based searches for social science research data discovery. arXiv preprint arXiv:2601.19559.",
    "Herbert, B., Szarvas, G., & Gurevych, I. (2009, September). Prior art search using international patent classification codes and all-claims-queries. In Workshop of the Cross-Language Evaluation Forum for European Languages (pp. 452-459). Springer.",
    "Islam Erana, T., & Finlayson, M. A. (2024). Automated neural patent landscaping in the small data regime. arXiv preprint arXiv:2407.08001.",
    "Madani, F., & Weber, C. (2016). The evolution of patent mining: Applying bibliometrics analysis and keyword network analysis. World Patent Information, 46, 32-48.",
    "Ratner, A. J., Bach, S. H., Ehrenberg, H., Fries, J., Wu, S., & Re, C. (2017). Snorkel: Rapid training data creation with weak supervision. Proceedings of the VLDB Endowment, 11(3), 269-282.",
    "Sofean, M. (2026). Identification of domain-relevant patents via weakly supervised deep learning. World Patent Information, 84, 102434.",
    "Trippe, A. (2015). Guidelines for preparing patent landscape reports. Patent landscape reports. Geneva: WIPO.",
    "Tseng, Y. H., Lin, C. J., & Lin, Y. I. (2007). Text mining techniques for patent analysis. Information Processing & Management, 43(5), 1216-1247.",
    "van Rijn, T., & Timmis, J. K. (2023). Patent landscape analysis: Contributing to the identification of technology trends and informing research and innovation funding policy. Microbial Biotechnology, 16(4), 683-696.",
    "Wang, L., Ma, C., Feng, X., Zhang, Z., Yang, H., Zhang, J., et al. (2024). A survey on large language model based autonomous agents. Frontiers of Computer Science, 18(6), 186345.",
    "Zaini, W. M. F., Lai, D. T. C., & Lim, R. C. (2022). Identifying patent classification codes associated with specific search keywords using machine learning. World Patent Information, 71, 102153.",
]


def display_text(text: str) -> str:
    replacements = {
        "Bergeaud and Verluise": "Bergeaud et al.",
        "allenai/scibert_scivocab_uncased": "allenai/scibert-scivocab-uncased",
        "abstain_candidate": "abstain candidate",
        "easy_positive": "easy positive",
        "easy_negative": "easy negative",
        "hard_negative": "hard negative",
        "candidate_type": "candidate type",
        "NOT_SEED": "NOT-SEED",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def abstract_text() -> str:
    text = display_text(ABSTRACT)
    return text.replace(
        "We evaluate the framework on six technology domains from Bergeaud et al.:",
        "We evaluate the framework on six technology domains from Bergeaud et al. (2017):",
    )


def cited_text(text: str) -> str:
    text = display_text(text)

    if text.startswith("Patent landscaping analysis provides"):
        text = text.replace(
            "surrounding a scientific or technological domain.",
            "surrounding a scientific or technological domain (Trippe, 2015; van Rijn & Timmis, 2023).",
        )

    elif text.startswith("Traditional patent identification relies"):
        text = text.replace(
            "expert-crafted Boolean queries.",
            "expert-crafted Boolean queries (Herbert et al., 2009; Tseng et al., 2007; Zaini et al., 2022).",
        )
        text = text.replace(
            "limited precision.",
            "limited precision (Madani & Weber, 2016; Cypris, 2026).",
        )

    elif text.startswith("Weak supervision addresses"):
        text = text.replace(
            "manual annotation.",
            "manual annotation (Ratner et al., 2017).",
        )
        text = text.replace(
            "Sofean's patent identification pipeline uses Snorkel",
            "Sofean's patent identification pipeline uses Snorkel (Sofean, 2026)",
        )

    elif text.startswith("This paper asks whether"):
        text = (
            text
            + " This design is also motivated by recent evidence that large language models can improve semantic search and autonomous reasoning workflows when their outputs are constrained by task-specific structure (Green et al., 2026; Wang et al., 2024)."
        )

    elif text.startswith("Patent landscaping and patent identification."):
        text = text.replace(
            "competitive landscapes.",
            "competitive landscapes (Trippe, 2015; van Rijn & Timmis, 2023).",
        )
        text = text.replace(
            "adjacent technologies.",
            "adjacent technologies (Herbert et al., 2009; Madani & Weber, 2016; Tseng et al., 2007; Zaini et al., 2022).",
        )

    elif text.startswith("Automated patent landscaping."):
        text = text.replace(
            "representative seed patents.",
            "representative seed patents (Abood & Feltenberger, 2018).",
        )
        text = text.replace(
            "seed and anti-seed benchmarks.",
            "seed and anti-seed benchmarks (Bergeaud et al., 2017).",
        )
        text += (
            " Later neural and transformer-based patent landscaping studies further show that semantic representations can support automated patent-set construction under limited-label conditions (Choi et al., 2019; Islam Erana & Finlayson, 2024)."
        )

    elif text.startswith("Weak supervision and Snorkel."):
        text = text.replace(
            "labeling functions.",
            "labeling functions (Ratner et al., 2017).",
            1,
        )
        text = text.replace(
            "static keyword rules.",
            "static keyword rules (Sofean, 2026).",
        )

    elif text.startswith("Large language models and agentic labeling."):
        text = text.replace(
            "natural language descriptions,",
            "natural language descriptions (Green et al., 2026),",
        )
        text = text.replace(
            "constrained multi-agent design instead:",
            "constrained multi-agent design instead (Wang et al., 2024):",
        )

    elif text.startswith("Data and domains."):
        text = text.replace(
            "drawn from Bergeaud et al.:",
            "drawn from Bergeaud et al. (2017):",
        )
        text += " The blockchain domain also reflects the broader patent-landscaping challenge that expert search queries can be transparent but labor-intensive to maintain (Clarke et al., 2020)."

    elif text.startswith("Snorkel baseline."):
        text = text.replace(
            "uses labeling functions and a LabelModel.",
            "uses labeling functions and a LabelModel, following the weak-supervision strategy used in Snorkel and domain-relevant patent identification (Ratner et al., 2017; Sofean, 2026).",
        )

    elif text.startswith("MAS labeler."):
        text = text.replace(
            "rather than hand-written labeling functions.",
            "rather than hand-written labeling functions, using agentic decomposition to separate relevance, exclusion, and deterministic scoring (Wang et al., 2024).",
        )

    elif text.startswith("Downstream training."):
        text = text.replace(
            "allenai/scibert-scivocab-uncased",
            "allenai/scibert-scivocab-uncased (Beltagy et al., 2019)",
        )

    elif text.startswith("Evaluation."):
        text = text.replace(
            "a technological field.",
            "a technological field (Trippe, 2015; van Rijn & Timmis, 2023).",
        )

    elif text.startswith("The experiment covers"):
        text = text.replace(
            "six technologies.",
            "six technologies defined in the semantic patent benchmark (Bergeaud et al., 2017).",
        )

    elif text.startswith("The first diagnostic result concerns"):
        text = text.replace(
            "keyword and CPC classes.",
            "keyword and CPC classes, a known limitation of search-query and classification-code approaches (Herbert et al., 2009; Zaini et al., 2022).",
        )

    elif text.startswith("The largest gains appear"):
        text = text.replace(
            "one-sided training signal.",
            "one-sided training signal, which is consistent with prior concerns about keyword-driven patent mining and expert search-query construction (Clarke et al., 2020; Madani & Weber, 2016).",
        )

    elif text.startswith("The results support"):
        text = text.replace(
            "patent identification.",
            "patent identification (Sofean, 2026).",
            1,
        )

    elif text.startswith("The most important mechanism"):
        text = text.replace(
            "functional-application level.",
            "functional-application level (Bergeaud et al., 2017).",
        )

    elif text.startswith("The framework also generalizes"):
        text = text.replace(
            "across domains.",
            "across domains, complementing prior work on automated and neural patent landscaping across technology areas (Abood & Feltenberger, 2018; Choi et al., 2019; Islam Erana & Finlayson, 2024).",
        )

    elif text.startswith("Several limitations remain"):
        text = text.replace(
            "LLM calls,",
            "LLM calls,",
        )
        text = text.replace(
            "model behavior.",
            "model behavior, an issue also noted in broader surveys of LLM-based agents (Wang et al., 2024).",
        )

    elif text.startswith("This paper proposed"):
        text = text.replace(
            "domain-relevant patent identification.",
            "domain-relevant patent identification, building on weakly supervised patent identification and automated patent-landscaping research (Abood & Feltenberger, 2018; Sofean, 2026).",
        )

    elif text.startswith("The central empirical finding"):
        text = text.replace(
            "high:",
            "high in patent landscape analysis:",
        )
        text += " This conclusion is consistent with recent practitioner discussions of AI-assisted patent landscape analysis, but the present study keeps the comparison experimentally controlled rather than treating LLMs as an unstructured search substitute (Cypris, 2026; Green et al., 2026)."

    return text


def set_run_font(run, size: float | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:cs"), FONT_NAME)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)


def set_style_font(style, size: float, bold: bool = False, italic: bool = False) -> None:
    style.font.name = FONT_NAME
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    style._element.rPr.rFonts.set(qn("w:cs"), FONT_NAME)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic
    style.font.color.rgb = RGBColor(0, 0, 0)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    set_run_font(run, 10)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def set_cell_borders(cell, **kwargs) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        if edge_data is None:
            element.set(qn("w:val"), "nil")
            continue
        for key, value in edge_data.items():
            element.set(qn(f"w:{key}"), str(value))


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[int]) -> None:
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.different_first_page_header_footer = True

    add_page_number(section.footer.paragraphs[0])
    add_page_number(section.first_page_footer.paragraphs[0])

    normal = doc.styles["Normal"]
    set_style_font(normal, 10.5)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(9)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, before, after in [
        ("Heading 1", 17, 16, 12),
        ("Heading 2", 13, 13, 8),
        ("Heading 3", 11, 10, 6),
    ]:
        style = doc.styles[style_name]
        set_style_font(style, size, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0


def add_heading_text(doc: Document, text: str, level: int = 1) -> None:
    if level == 1:
        size, before, after = 17, 16, 12
    elif level == 2:
        size, before, after = 13, 13, 8
    else:
        size, before, after = 11, 10, 6
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(display_text(text))
    set_run_font(r, size, bold=True)


def add_title_block(doc: Document) -> None:
    doc.add_paragraph().paragraph_format.space_after = Pt(44)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    r = p.add_run(TITLE)
    set_run_font(r, 20)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(AUTHORS)
    set_run_font(r, 12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    r = p.add_run(AFFILIATION)
    set_run_font(r, 11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Abstract")
    set_run_font(r, 12, bold=True)

    add_body_paragraph(doc, abstract_text(), after=15, cited=False)


def add_body_paragraph(doc: Document, text: str, after: float = 9, cited: bool = True) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(cited_text(text) if cited else display_text(text))
    set_run_font(r, 10.5)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(display_text(text))
    set_run_font(r, 10)


def add_figure(doc: Document, image_path: Path, caption: str) -> None:
    if image_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(image_path), width=Inches(6.45))
    add_caption(doc, caption)


def add_booktabs_table(
    doc: Document,
    caption: str,
    headers: list[str],
    rows: list[list[str]],
    widths: list[int],
) -> None:
    add_caption(doc, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, widths)

    header_line = {"val": "single", "sz": "8", "space": "0", "color": "000000"}
    mid_line = {"val": "single", "sz": "4", "space": "0", "color": "000000"}
    bottom_line = {"val": "single", "sz": "8", "space": "0", "color": "000000"}

    header_cells = table.rows[0].cells
    for idx, value in enumerate(headers):
        cell = header_cells[idx]
        cell.text = ""
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(display_text(value))
        set_run_font(r, 9.4)
        set_cell_borders(cell, top=header_line, bottom=mid_line)

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            cell.text = ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            r = p.add_run(display_text(str(value)))
            set_run_font(r, 9.4)
            set_cell_borders(cell)

    for cell in table.rows[-1].cells:
        set_cell_borders(cell, bottom=bottom_line)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)


def add_results_tables(doc: Document) -> None:
    add_heading_text(doc, "4.1 Dataset Scale", level=2)
    add_booktabs_table(
        doc,
        "Table 1: Gold evaluation sets and in-domain labeling candidate pools.",
        ["Domain", "SEED", "NOT-SEED", "Gold N", "Labeling Pool N"],
        [[r[0], f"{r[1]:,}", f"{r[2]:,}", f"{r[3]:,}", f"{r[4]:,}"] for r in DATASET_ROWS],
        [3300, 1050, 1250, 1250, 2510],
    )

    add_heading_text(doc, "4.2 Labeling Diagnostics", level=2)
    add_booktabs_table(
        doc,
        "Table 2: In-domain label distribution and OOD false positives. MAS negatives combine easy negative and hard negative; Snorkel negatives are NOT-SEED labels.",
        ["Domain", "Labeler", "Positive", "Negative", "Boundary / Abstain", "OOD FP"],
        [[r[0], r[1], f"{r[2]:,}", f"{r[3]:,}", f"{r[4]:,}", f"{r[5]:,}"] for r in LABEL_DIAGNOSTICS],
        [3000, 1050, 1150, 1150, 1850, 1160],
    )

    add_heading_text(doc, "4.3 Downstream Performance", level=2)
    add_booktabs_table(
        doc,
        "Table 3: SciBERT evaluation on held-out gold benchmarks at threshold 0.5.",
        ["Domain", "Labeler", "AUC", "Macro-F1", "Recall", "Precision", "Accuracy"],
        [
            [r[0], r[1], f"{r[2]:.3f}", f"{r[3]:.3f}", f"{r[4]:.3f}", f"{r[5]:.3f}", f"{r[6]:.3f}"]
            for r in PERFORMANCE_ROWS
        ],
        [2700, 1050, 850, 1050, 900, 1050, 1100],
    )


def add_references(doc: Document) -> None:
    add_heading_text(doc, "References", level=1)
    for ref in FINAL_REFERENCES:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(display_text(ref))
        set_run_font(r, 10)


def guid_key_bytes(guid: uuid.UUID) -> bytes:
    raw = guid.bytes
    return raw[3::-1] + raw[5:3:-1] + raw[7:5:-1] + raw[8:]


def obfuscate_font(font_path: Path, key: uuid.UUID) -> bytes:
    data = bytearray(font_path.read_bytes())
    key_bytes = guid_key_bytes(key)
    for idx in range(min(32, len(data))):
        data[idx] ^= key_bytes[idx % 16]
    return bytes(data)


def ensure_xml_part(zf: zipfile.ZipFile, name: str, fallback: bytes) -> bytes:
    try:
        return zf.read(name)
    except KeyError:
        return fallback


def embed_cmu_fonts(docx_path: Path) -> None:
    font_files = {
        "embedRegular": FONT_DIR / "cmunrm.ttf",
        "embedBold": FONT_DIR / "cmunbx.ttf",
        "embedItalic": FONT_DIR / "cmunti.ttf",
        "embedBoldItalic": FONT_DIR / "cmunbi.ttf",
    }
    if not all(path.exists() for path in font_files.values()):
        return

    ns = {
        "ct": "http://schemas.openxmlformats.org/package/2006/content-types",
        "rel": "http://schemas.openxmlformats.org/package/2006/relationships",
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    }
    for prefix, uri in ns.items():
        ET.register_namespace("" if prefix in {"ct", "rel"} else prefix, uri)

    with tempfile.TemporaryDirectory() as tmp:
        tmp_docx = Path(tmp) / docx_path.name
        with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp_docx, "w", zipfile.ZIP_DEFLATED) as zout:
            existing = set(zin.namelist())

            content_types = ET.fromstring(zin.read("[Content_Types].xml"))
            if not any(child.attrib.get("Extension") == "odttf" for child in content_types):
                default = ET.Element(f"{{{ns['ct']}}}Default")
                default.set("Extension", "odttf")
                default.set("ContentType", "application/vnd.openxmlformats-officedocument.obfuscatedFont")
                content_types.append(default)

            font_table_xml = ensure_xml_part(
                zin,
                "word/fontTable.xml",
                f'<w:fonts xmlns:w="{ns["w"]}" xmlns:r="{ns["r"]}"/>'.encode("utf-8"),
            )
            font_table = ET.fromstring(font_table_xml)
            for font in list(font_table.findall("w:font", ns)):
                if font.attrib.get(f"{{{ns['w']}}}name") == FONT_NAME:
                    font_table.remove(font)

            font_node = ET.Element(f"{{{ns['w']}}}font")
            font_node.set(f"{{{ns['w']}}}name", FONT_NAME)
            family = ET.SubElement(font_node, f"{{{ns['w']}}}family")
            family.set(f"{{{ns['w']}}}val", "roman")
            pitch = ET.SubElement(font_node, f"{{{ns['w']}}}pitch")
            pitch.set(f"{{{ns['w']}}}val", "variable")

            rels_xml = ensure_xml_part(
                zin,
                "word/_rels/fontTable.xml.rels",
                f'<Relationships xmlns="{ns["rel"]}"/>'.encode("utf-8"),
            )
            rels = ET.fromstring(rels_xml)
            for rel in list(rels):
                if rel.attrib.get("Target", "").startswith("fonts/cmu"):
                    rels.remove(rel)

            embedded_parts: dict[str, bytes] = {}
            for idx, (tag, font_path) in enumerate(font_files.items(), start=1):
                rel_id = f"rIdCMU{idx}"
                key = uuid.uuid5(uuid.NAMESPACE_URL, f"{FONT_NAME}-{tag}")
                embed = ET.SubElement(font_node, f"{{{ns['w']}}}{tag}")
                embed.set(f"{{{ns['r']}}}id", rel_id)
                embed.set(f"{{{ns['w']}}}fontKey", "{" + str(key).upper() + "}")
                embed.set(f"{{{ns['w']}}}subsetted", "false")

                target = f"fonts/cmu{idx}.odttf"
                rel = ET.Element(f"{{{ns['rel']}}}Relationship")
                rel.set("Id", rel_id)
                rel.set("Type", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/font")
                rel.set("Target", target)
                rels.append(rel)
                embedded_parts[f"word/{target}"] = obfuscate_font(font_path, key)

            font_table.append(font_node)

            settings_xml = zin.read("word/settings.xml")
            settings = ET.fromstring(settings_xml)
            if settings.find("w:embedTrueTypeFonts", ns) is None:
                settings.append(ET.Element(f"{{{ns['w']}}}embedTrueTypeFonts"))
            save_subset = settings.find("w:saveSubsetFonts", ns)
            if save_subset is None:
                save_subset = ET.Element(f"{{{ns['w']}}}saveSubsetFonts")
                settings.append(save_subset)
            save_subset.set(f"{{{ns['w']}}}val", "false")

            replacements = {
                "[Content_Types].xml": ET.tostring(content_types, encoding="utf-8", xml_declaration=True),
                "word/fontTable.xml": ET.tostring(font_table, encoding="utf-8", xml_declaration=True),
                "word/_rels/fontTable.xml.rels": ET.tostring(rels, encoding="utf-8", xml_declaration=True),
                "word/settings.xml": ET.tostring(settings, encoding="utf-8", xml_declaration=True),
            }

            for item in zin.infolist():
                if item.filename in replacements or item.filename in embedded_parts:
                    continue
                zout.writestr(item, zin.read(item.filename))
            for name, data in replacements.items():
                zout.writestr(name, data)
                existing.add(name)
            for name, data in embedded_parts.items():
                zout.writestr(name, data)

        shutil.move(str(tmp_docx), docx_path)


def build_docx() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_doc(doc)
    add_title_block(doc)

    for section_idx, (heading, paragraphs) in enumerate(SECTIONS[:3]):
        add_heading_text(doc, heading, level=1)
        for paragraph in paragraphs:
            add_body_paragraph(doc, paragraph)
        if section_idx == 0:
            add_figure(
                doc,
                FIG_OVERALL,
                "Figure 1: Overall controlled framework for comparing Snorkel and MAS weak supervision in domain-relevant patent identification.",
            )
        if section_idx == 2:
            add_figure(
                doc,
                FIG_MAS,
                "Figure 2: MAS labeling framework. Node A extracts relevance evidence and routes the patent; Node B checks boundary and hard-negative exclusions; Node C deterministically maps the structured state to a pseudo-label.",
            )

    heading, paragraphs = SECTIONS[3]
    add_heading_text(doc, heading, level=1)
    for paragraph in paragraphs:
        add_body_paragraph(doc, paragraph)
    add_results_tables(doc)

    for heading, paragraphs in SECTIONS[4:]:
        add_heading_text(doc, heading, level=1)
        for paragraph in paragraphs:
            add_body_paragraph(doc, paragraph)

    add_references(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_docx()
